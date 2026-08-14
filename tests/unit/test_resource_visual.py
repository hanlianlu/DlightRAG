# Copyright 2025-2026 Hanlian Lu. SPDX-License-Identifier: Apache-2.0
"""Tests for focused visual resource inspection through the VLM role."""

from __future__ import annotations

import io
import threading
from typing import Any

import openpyxl
import pytest
from dlightrag_ai.media import decode_image_base64
from docx import Document
from openpyxl.drawing.image import Image as XLImage
from PIL import Image

from dlightrag.core.resources.models import ResourceInput
from dlightrag.core.resources.registry import ResourceRegistry
from dlightrag.core.resources.visual import (
    ResourceInspectionError,
    ResourceInspector,
)
from tests.unit.conftest import answer_image_policy

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _inspector(
    registry: ResourceRegistry, vlm: Any, *, max_images: int = 8, **overrides: Any
) -> ResourceInspector:
    return ResourceInspector(
        registry,
        vlm_func=vlm,
        image_policy=answer_image_policy(max_images=max_images),
        **overrides,
    )


class _RecordingVLM:
    """Stub VLM callable that records the messages it receives."""

    def __init__(self, reply: str = "The figure shows rising Q3 revenue.") -> None:
        self.reply = reply
        self.calls: list[list[dict]] = []

    async def __call__(self, *, messages: list[dict], **_kwargs) -> str:
        self.calls.append(messages)
        return self.reply


class _FailingVLM:
    async def __call__(self, *, messages: list[dict], **_kwargs) -> str:
        raise RuntimeError("vlm upstream 503")


def _png(color: tuple[int, int, int], size: tuple[int, int] = (24, 24)) -> bytes:
    buffer = io.BytesIO()
    Image.new("RGB", size, color).save(buffer, "PNG")
    return buffer.getvalue()


def _pdf_bytes(pages: list[tuple[int, int]]) -> bytes:
    images = [Image.new("RGB", size, (10, 10, 10)) for size in pages]
    buffer = io.BytesIO()
    images[0].save(buffer, "PDF", save_all=True, append_images=images[1:])
    return buffer.getvalue()


def _docx_bytes(image: bytes) -> bytes:
    document = Document()
    document.add_heading("Quarterly Report", level=1)
    document.add_paragraph("Revenue trends are summarized below.")
    document.add_picture(io.BytesIO(image))
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _xlsx_bytes(image: bytes) -> bytes:
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    assert worksheet is not None
    worksheet.title = "Financials"
    worksheet["A1"] = "Revenue"
    worksheet.add_image(XLImage(io.BytesIO(image)), "B2")
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _sent_image_sizes(messages: list[dict]) -> list[tuple[int, int]]:
    content = messages[-1]["content"]
    sizes: list[tuple[int, int]] = []
    for block in content:
        if block.get("type") != "image_url":
            continue
        raw, _ = decode_image_base64(block["image_url"]["url"])
        with Image.open(io.BytesIO(raw)) as image:
            sizes.append(image.size)
    return sizes


async def test_inspect_source_image_returns_vlm_evidence() -> None:
    vlm = _RecordingVLM("A bar chart with three ascending bars.")
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(
                filename="chart.png", content=_png((200, 30, 30)), declared_mime="image/png"
            )
        )
        inspector = _inspector(registry, vlm)

        result = await inspector.inspect(resource_id, "What does the chart show?")

    assert result.derived_by_vlm is True
    assert result.content == "A bar chart with three ascending bars."
    assert result.locator.kind == "image"
    assert len(vlm.calls) == 1
    assert len(_sent_image_sizes(vlm.calls[0])) == 1


async def test_inspection_image_budgeting_runs_off_event_loop(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from dlightrag_ai.media import ImagePayloadBudget

    loop_thread = threading.get_ident()
    budget_threads: list[int] = []
    add_base64 = ImagePayloadBudget.add_base64

    def capture_budget(self, value, *, label):
        budget_threads.append(threading.get_ident())
        return add_base64(self, value, label=label)

    monkeypatch.setattr(ImagePayloadBudget, "add_base64", capture_budget)
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(
                filename="chart.png",
                content=_png((200, 30, 30)),
                declared_mime="image/png",
            )
        )
        inspector = _inspector(registry, _RecordingVLM())

        await inspector.inspect(resource_id, "Read the chart")

    assert budget_threads
    assert all(thread_id != loop_thread for thread_id in budget_threads)


async def test_inspect_pdf_overview_is_bounded_and_low_resolution() -> None:
    vlm = _RecordingVLM()
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(filename="deck.pdf", content=_pdf_bytes([(300, 400)] * 3))
        )
        inspector = _inspector(registry, vlm)

        result = await inspector.inspect(resource_id, "Which page has the revenue table?")

    assert result.locator.kind == "pdf_overview"
    assert result.locator.page_start == 1
    assert result.locator.page_end == 3
    assert result.has_more is False
    overview_sizes = _sent_image_sizes(vlm.calls[0])
    assert len(overview_sizes) == 3
    assert max(max(size) for size in overview_sizes) <= 400


async def test_inspect_pdf_selected_full_page_is_higher_resolution() -> None:
    vlm = _RecordingVLM()
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(filename="deck.pdf", content=_pdf_bytes([(300, 400)] * 3))
        )
        inspector = _inspector(registry, vlm)

        result = await inspector.inspect(resource_id, "Read the table", locator="2")

    assert result.locator.kind == "pdf_page"
    assert result.locator.page == 2
    page_sizes = _sent_image_sizes(vlm.calls[0])
    assert len(page_sizes) == 1
    assert max(page_sizes[0]) >= 600


async def test_inspect_pdf_missing_page_raises() -> None:
    vlm = _RecordingVLM()
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(filename="deck.pdf", content=_pdf_bytes([(300, 400)]))
        )
        inspector = _inspector(registry, vlm)

        with pytest.raises(ResourceInspectionError):
            await inspector.inspect(resource_id, "Read page 9", locator="9")
    assert vlm.calls == []


async def test_inspect_rejects_invalid_or_conflicting_locators() -> None:
    vlm = _RecordingVLM()
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(filename="deck.pdf", content=_pdf_bytes([(300, 400)]))
        )
        inspector = _inspector(registry, vlm)

        with pytest.raises(ResourceInspectionError, match="page number"):
            await inspector.inspect(resource_id, "Read the page", locator="page nope")
        with pytest.raises(ResourceInspectionError, match="mutually exclusive"):
            await inspector.inspect(resource_id, "Read the page", locator="1", cursor="cursor")

    assert vlm.calls == []


async def test_inspect_source_image_rejects_page_cursor() -> None:
    vlm = _RecordingVLM()
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(filename="chart.png", content=_png((2, 3, 4)), declared_mime="image/png")
        )
        inspector = _inspector(registry, vlm)

        with pytest.raises(ResourceInspectionError, match="source image"):
            await inspector.inspect(resource_id, "Read the chart", cursor="cursor")

    assert vlm.calls == []


async def test_inspect_pdf_overview_paginates() -> None:
    vlm = _RecordingVLM()
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(filename="deck.pdf", content=_pdf_bytes([(200, 260)] * 3))
        )
        inspector = _inspector(registry, vlm, overview_page_limit=2)

        first = await inspector.inspect(resource_id, "Find the chart")
        assert first.locator.page_start == 1
        assert first.locator.page_end == 2
        assert first.has_more is True
        assert first.next_cursor is not None

        second = await inspector.inspect(resource_id, "Find the chart", cursor=first.next_cursor)
    assert second.locator.page_start == 3
    assert second.locator.page_end == 3
    assert second.has_more is False


async def test_inspect_pdf_overview_locator_only_claims_pages_sent_to_vlm() -> None:
    vlm = _RecordingVLM()
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(filename="deck.pdf", content=_pdf_bytes([(200, 260)] * 3))
        )
        inspector = _inspector(
            registry,
            vlm,
            max_images=1,
            overview_page_limit=3,
        )

        first = await inspector.inspect(resource_id, "Find the chart")
        assert first.locator.page_start == first.locator.page_end == 1
        assert first.next_cursor is not None
        second = await inspector.inspect(resource_id, "Find the chart", cursor=first.next_cursor)

    assert second.locator.page_start == second.locator.page_end == 2


async def test_inspect_docx_embedded_visual_handle() -> None:
    vlm = _RecordingVLM("A red company logo.")
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(
                filename="report.docx",
                content=_docx_bytes(_png((220, 20, 20))),
                declared_mime=DOCX_MIME,
            )
        )
        read = await registry.read(resource_id)
        handle_id = read.visual_handles[0].handle_id
        inspector = _inspector(registry, vlm)

        result = await inspector.inspect(resource_id, "Describe the image", locator=handle_id)

    assert result.locator.kind == "visual"
    assert result.locator.handle_id == handle_id
    assert result.derived_by_vlm is True
    assert result.content == "A red company logo."


async def test_inspect_xlsx_sheet_cell_handle_carries_anchor() -> None:
    vlm = _RecordingVLM("A pie chart.")
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(
                filename="model.xlsx",
                content=_xlsx_bytes(_png((10, 10, 220))),
                declared_mime=XLSX_MIME,
            )
        )
        read = await registry.read(resource_id)
        handle_id = read.visual_handles[0].handle_id
        inspector = _inspector(registry, vlm)

        result = await inspector.inspect(resource_id, "Describe the image", locator=handle_id)

    assert result.locator.kind == "visual"
    assert result.locator.anchor == "Financials!B2"


async def test_inspect_vlm_failure_surfaces_as_inspection_error() -> None:
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(filename="chart.png", content=_png((1, 2, 3)), declared_mime="image/png")
        )
        inspector = _inspector(registry, _FailingVLM())

        with pytest.raises(ResourceInspectionError, match="visual inspection failed") as failure:
            await inspector.inspect(resource_id, "What is this?")

    assert "503" not in str(failure.value)


async def test_inspect_text_resource_without_handle_raises() -> None:
    vlm = _RecordingVLM()
    async with ResourceRegistry() as registry:
        resource_id = registry.register(
            ResourceInput(filename="notes.txt", content=b"just text", declared_mime="text/plain")
        )
        inspector = _inspector(registry, vlm)

        with pytest.raises(ResourceInspectionError):
            await inspector.inspect(resource_id, "Look at it")
    assert vlm.calls == []
