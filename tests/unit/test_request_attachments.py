# Copyright 2025-2026 Hanlian Lu. SPDX-License-Identifier: Apache-2.0
"""Tests for Web Composer document parsing and budget packing."""

import asyncio
import io
import json
import threading
from contextlib import asynccontextmanager
from dataclasses import replace
from pathlib import Path
from types import SimpleNamespace
from typing import Any
from unittest.mock import AsyncMock

import numpy as np
import pytest
from lightrag.utils import TiktokenTokenizer
from PIL import Image

from dlightrag.core.request import attachment_digest, attachments
from dlightrag.core.request.attachment_digest import (
    ATTACHMENT_PLANNER_DIGEST_TOKEN_BUDGET,
    build_attachment_planner_digests,
)
from dlightrag.core.request.attachments import (
    AttachmentContextChunk,
    ComposerDocumentService,
    ParsedAttachmentBundle,
    _ParseOwnerShim,
    parse_attachment_to_bundle,
    resolve_attachment_chunk_signature,
    resolve_attachment_parser_signature,
)
from dlightrag.utils.tokens import estimate_tokens


def text_chunk(chunk_id: str, tokens: int = 10) -> AttachmentContextChunk:
    return AttachmentContextChunk(
        chunk_id=chunk_id,
        attachment_id="att-1",
        filename="report.pdf",
        chunk_index=1,
        content="alpha",
        token_estimate=tokens,
    )


def visual_chunk(chunk_id: str) -> AttachmentContextChunk:
    return AttachmentContextChunk(
        chunk_id=chunk_id,
        attachment_id="att-1",
        filename="report.pdf",
        chunk_index=2,
        content="chart description",
        token_estimate=10,
        sidecar_type="drawing",
        image_bytes=b"png-bytes",
        image_mime_type="image/png",
    )


_RETRIEVAL_CONTENT = "retrieval evidence " * 30_000


def retrieval_text_chunk(chunk_id: str) -> AttachmentContextChunk:
    return replace(
        text_chunk(chunk_id),
        content=_RETRIEVAL_CONTENT,
        token_estimate=estimate_tokens(_RETRIEVAL_CONTENT),
    )


def retrieval_visual_chunk(chunk_id: str) -> AttachmentContextChunk:
    return replace(
        visual_chunk(chunk_id),
        content=_RETRIEVAL_CONTENT,
        token_estimate=estimate_tokens(_RETRIEVAL_CONTENT),
    )


def test_attachment_evidence_mode_routes_exact_default_budget_to_full() -> None:
    content = "x" * (24_576 * 4)
    chunks = [
        attachments.build_text_attachment_chunk(
            attachment_id="att-boundary-full",
            filename="boundary-full.txt",
            chunk_id="boundary-full",
            chunk_index=1,
            content=content,
        )
    ]

    estimated_tokens = sum(estimate_tokens(chunk.content) for chunk in chunks)

    assert estimated_tokens == 24_576
    assert sum(chunk.token_estimate for chunk in chunks) == estimated_tokens
    assert attachments._resolve_attachment_evidence_mode(chunks) == "full"


def test_attachment_evidence_mode_routes_one_token_over_default_budget_to_retrieval() -> None:
    content = "x" * (24_577 * 4)
    chunks = [
        attachments.build_text_attachment_chunk(
            attachment_id="att-boundary-retrieval",
            filename="boundary-retrieval.txt",
            chunk_id="boundary-retrieval",
            chunk_index=1,
            content=content,
        )
    ]

    estimated_tokens = sum(estimate_tokens(chunk.content) for chunk in chunks)

    assert estimated_tokens == 24_577
    assert sum(chunk.token_estimate for chunk in chunks) == estimated_tokens
    assert attachments._resolve_attachment_evidence_mode(chunks) == "retrieval"


def _png_bytes(*, size: tuple[int, int]) -> bytes:
    buffer = io.BytesIO()
    Image.new("RGB", size, "white").save(buffer, format="PNG")
    return buffer.getvalue()


def planner_bundle(
    attachment_id: str,
    contents: list[str],
    *,
    sidecar_types: dict[int, str] | None = None,
) -> ParsedAttachmentBundle:
    return ParsedAttachmentBundle(
        chunks=[
            AttachmentContextChunk(
                chunk_id=f"{attachment_id}-{index}",
                attachment_id=attachment_id,
                filename=f"{attachment_id}.pdf",
                chunk_index=index,
                content=content,
                token_estimate=estimate_tokens(content),
                sidecar_type=(sidecar_types or {}).get(index),
            )
            for index, content in enumerate(contents, start=1)
        ]
    )


def test_attachment_planner_digest_keeps_short_document_in_full() -> None:
    bundle = planner_bundle(
        "short",
        [
            "Fractions Challenge Worksheet\nInstructions: simplify every fraction.",
            "Problem 20: Explain how the numerator and denominator change.",
        ],
    )

    digests, trace = build_attachment_planner_digests([("short", bundle)])

    assert "Fractions Challenge Worksheet" in digests["short"]
    assert "Problem 20" in digests["short"]
    assert trace["attachment_digest_strategy"] == "full"
    assert trace["attachment_digest_output_tokens"] <= ATTACHMENT_PLANNER_DIGEST_TOKEN_BUDGET


def test_attachment_planner_digest_preserves_structure_and_uniform_coverage() -> None:
    contents = [f"Section {index} " + (f"body-{index} " * 700) for index in range(24)]
    contents[0] = "START-MARKER Document title and executive introduction. " + contents[0]
    contents[8] = "[Table Name] Revenue by region\nColumns: Region, Revenue. " + contents[8]
    contents[12] = "MIDDLE-MARKER Central findings. " + contents[12]
    contents[-1] = "END-MARKER Final conclusion and recommendations. " + contents[-1]
    bundle = planner_bundle("long", contents, sidecar_types={9: "table"})

    digests, trace = build_attachment_planner_digests([("long", bundle)])
    digest = digests["long"]

    assert "START-MARKER" in digest
    assert "[Table Name] Revenue by region" in digest
    assert "MIDDLE-MARKER" in digest
    assert "END-MARKER" in digest
    assert trace["attachment_digest_strategy"] == "sampled"
    assert trace["attachment_digest_output_tokens"] <= ATTACHMENT_PLANNER_DIGEST_TOKEN_BUDGET


def test_attachment_planner_digest_shares_global_budget_across_documents() -> None:
    bundles = [
        (
            f"doc-{index}",
            planner_bundle(
                f"doc-{index}",
                [f"DOC-{index}-SECTION-{part} " + ("detail " * 700) for part in range(16)],
            ),
        )
        for index in range(3)
    ]

    digests, trace = build_attachment_planner_digests(bundles)

    assert set(digests) == {"doc-0", "doc-1", "doc-2"}
    assert all(digests.values())
    budgets = trace["attachment_digest_document_budgets"]
    assert sum(budgets.values()) == ATTACHMENT_PLANNER_DIGEST_TOKEN_BUDGET
    assert min(budgets.values()) >= 1_536
    assert max(budgets.values()) - min(budgets.values()) <= 1
    assert trace["attachment_digest_output_tokens"] <= ATTACHMENT_PLANNER_DIGEST_TOKEN_BUDGET

    _, reversed_trace = build_attachment_planner_digests(list(reversed(bundles)))
    assert reversed_trace["attachment_digest_document_budgets"] == budgets


def test_attachment_planner_token_samples_do_not_cluster_at_the_front() -> None:
    contents = [f"chunk-{index} " + ("x " * 200) for index in range(24)]

    selected = attachment_digest._uniform_token_indices(contents, list(range(24)), 7)

    assert selected[0] == 0
    assert selected[-1] == 23
    assert any(abs(index - 12) <= 1 for index in selected)
    assert max(right - left for left, right in zip(selected, selected[1:], strict=False)) <= 6


def test_attachment_context_chunk_emits_image_data_and_source_metadata() -> None:
    row = visual_chunk("v1").to_context_row()

    assert row["reference_id"] == "att-1"
    assert row["full_doc_id"] == "att-1"
    assert row["_workspace"] == "__web_attachment__"
    assert row["metadata"]["source_uri"] == "web-attachment://att-1"
    assert row["metadata"]["source_download_locator"] == "web-attachment://att-1"
    assert row["image_data"] == "cG5nLWJ5dGVz"
    assert row["image_mime_type"] == "image/png"
    # No pre-budget flag: the single AnswerContextPacker owns image budgeting.
    assert "_answer_image_sent" not in row


def test_attachment_context_chunk_keeps_full_cache_identity_private() -> None:
    cache_key = attachments.AttachmentCacheKey(
        content_sha256="content-v1",
        parser_signature="parser-v1",
        chunk_signature="chunker-v1",
        cache_chunk_id="cache-chunk-1",
    )
    chunk = AttachmentContextChunk(
        chunk_id="citation-chunk-1",
        attachment_id="att-1",
        filename="report.pdf",
        chunk_index=1,
        content="alpha",
        cache_key=cache_key,
        embedding_signature="embed-v1",
        embedding_vector=[1.0, 0.0],
    )

    row = chunk.to_context_row()

    assert row["chunk_id"] == "citation-chunk-1"
    assert row["_cache_key"] == cache_key
    assert "embedding_signature" not in row
    assert "embedding_vector" not in row
    assert "_cache_key" not in row["metadata"]


def test_validate_attachment_vector_returns_matching_float32_vector() -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache-1")
    row = attachments.AttachmentVectorPageRow(
        global_order=7,
        cache_key=cache_key,
        embedding_signature="embed-v1",
        embedding_vector=[3, 4],
    )

    vector = attachments.validate_attachment_vector(
        row,
        expected_signature="embed-v1",
        expected_dimension=2,
    )

    assert vector is not None
    assert vector.dtype == np.float32
    np.testing.assert_array_equal(vector, np.asarray([3, 4], dtype=np.float32))


@pytest.mark.parametrize(
    ("signature", "vector", "expected_dimension"),
    [
        ("stale", [1.0, 2.0], 2),
        ("embed-v1", [1.0], 2),
        ("embed-v1", [float("nan"), 1.0], 2),
        ("embed-v1", [float("inf"), 1.0], 2),
        ("embed-v1", [0.0, 0.0], 2),
        ("embed-v1", ["not-a-number", 1.0], 2),
    ],
)
def test_validate_attachment_vector_treats_invalid_rows_as_misses(
    signature: str,
    vector: list[object],
    expected_dimension: int,
) -> None:
    row = attachments.AttachmentVectorPageRow(
        global_order=0,
        cache_key=attachments.AttachmentCacheKey("content", "parser", "chunker", "cache-1"),
        embedding_signature=signature,
        embedding_vector=vector,
    )

    assert (
        attachments.validate_attachment_vector(
            row,
            expected_signature="embed-v1",
            expected_dimension=expected_dimension,
        )
        is None
    )


def test_embedding_signature_uses_resolved_embedder_contract(test_config: Any) -> None:
    embedder: Any = SimpleNamespace(
        dimension=1024,
        image_enabled=True,
        asymmetric=True,
        min_image_pixel=32,
    )

    payload = json.loads(
        attachments.build_composer_embedding_signature(
            config=test_config,
            embedder=embedder,
            mode="fused",
        )
    )

    assert payload["contract_version"] == 2
    assert payload["asymmetric"] is True
    assert payload["image_normalization"] == {"min_image_pixel": 32}


async def test_dense_rankings_embed_query_once_without_embedding_documents(
    test_config: Any,
) -> None:
    current_key = attachments.AttachmentCacheKey(
        "content-v1",
        "parser-v1",
        "chunker-v1",
        "current",
    )
    history_key = attachments.AttachmentCacheKey(
        "content-v1",
        "parser-v1",
        "chunker-v1",
        "history",
    )
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        asymmetric=True,
        min_image_pixel=32,
        aembed_query=AsyncMock(return_value=[1.0, 0.0]),
        aembed_documents=AsyncMock(side_effect=AssertionError("documents must stay cached")),
    )
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    store = _SpyStore(
        None,
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                global_order=0,
                cache_key=current_key,
                embedding_signature=signature,
                embedding_vector=[0.6, 0.8],
            ),
            attachments.AttachmentVectorPageRow(
                global_order=1,
                cache_key=history_key,
                embedding_signature=signature,
                embedding_vector=[1.0, 0.0],
            ),
        ],
    )
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
        principal_id="p1",
        conversation_id="c1",
    )
    current_row = {
        "chunk_id": "current",
        "reference_id": "current-att",
        "content": "lower cosine current evidence",
        "_cache_key": current_key,
        "embedding_signature": signature,
        "embedding_vector": [0.6, 0.8],
        "metadata": {},
    }
    history_row = {
        "chunk_id": "history",
        "reference_id": "history-att",
        "content": "higher cosine history evidence",
        "_cache_key": history_key,
        "embedding_signature": signature,
        "embedding_vector": [1.0, 0.0],
        "metadata": {},
    }

    ranked, trace = await service.adense_rankings(
        "semantic query",
        [current_row, history_row],
    )

    assert [item["chunk_id"] for item in ranked] == ["history", "current"]
    assert trace == {
        "composer_dense_status": "ranked",
        "composer_dense_chunks": 2,
    }
    embedder.aembed_query.assert_awaited_once_with("semantic query")
    embedder.aembed_documents.assert_not_awaited()
    assert store.vector_calls == [
        {
            "principal_id": "p1",
            "conversation_id": "c1",
            "references": [(0, current_key), (1, history_key)],
            "ttl_days": 30,
            "page_size": 256,
        }
    ]
    assert all("embedding_signature" not in row for row in ranked)
    assert all("embedding_vector" not in row for row in ranked)


class _FakeLightRAG:
    """A minimal object exposing only the read-only parse/chunk knobs.

    Deliberately carries NO storage handles: reusing it for parsing must never
    reach a workspace store.
    """

    def __init__(self) -> None:
        self.tokenizer = TiktokenTokenizer()
        self.addon_params: dict[str, Any] = {}
        self.chunk_token_size = 1200
        self.embedding_func = None

    def _build_mm_chunks_from_sidecars(self, **_kwargs: Any) -> list[dict[str, Any]]:
        return []


async def test_bundle_uses_pre_rendered_mm_chunks_without_bound_renderer(tmp_path: Path) -> None:
    class _NoRendererLightRAG(_FakeLightRAG):
        def _build_mm_chunks_from_sidecars(self, **_kwargs: Any) -> list[dict[str, Any]]:
            raise AssertionError("workspace renderer must not be called")

    blocks_path = tmp_path / "report.blocks.jsonl"
    blocks_path.write_text('{"type":"meta"}\n', encoding="utf-8")
    rendered_mm_chunks = (
        {
            "chunk_id": "mm-drawing-000",
            "chunk_order_index": 0,
            "content": "analyzed chart",
            "sidecar": {"type": "drawing", "id": "chart-1"},
        },
    )

    bundle = await attachments.build_attachment_bundle_from_parse_result(
        lightrag=_NoRendererLightRAG(),
        attachment_id="att-1",
        filename="report.pdf",
        parsed=attachments.ParsedAttachmentDocument(
            content="ordinary parser text",
            blocks_path=str(blocks_path),
            parser_signature="legacy:",
        ),
        process_options="iP",
        multimodal_chunks=rendered_mm_chunks,
    )

    assert [chunk.content for chunk in bundle.chunks] == [
        "ordinary parser text",
        "analyzed chart",
    ]
    assert "_composer_mm_rendered" not in bundle.chunks[0].metadata
    assert bundle.chunks[1].metadata["_composer_mm_rendered"] is True
    assert bundle.chunks[1].sidecar_type == "drawing"


class _SpyStore:
    def __init__(
        self,
        cached: ParsedAttachmentBundle | None,
        *,
        vector_rows: list[attachments.AttachmentVectorPageRow] | None = None,
    ) -> None:
        self._cached = cached
        self._vector_rows = vector_rows or []
        self.load_calls: list[dict[str, Any]] = []
        self.materialized: list[ParsedAttachmentBundle] = []
        self.updated: list[list[AttachmentContextChunk]] = []
        self.vector_calls: list[dict[str, Any]] = []

    async def load_attachment_chunks(
        self,
        principal_id: str,
        conversation_id: str,
        attachment_id: str,
        filename: str,
        *,
        content_sha256: str,
        parser_signature: str,
        chunk_signature: str,
        ttl_days: int,
    ) -> ParsedAttachmentBundle | None:
        self.load_calls.append(
            {
                "principal_id": principal_id,
                "conversation_id": conversation_id,
                "attachment_id": attachment_id,
                "content_sha256": content_sha256,
                "parser_signature": parser_signature,
                "chunk_signature": chunk_signature,
                "ttl_days": ttl_days,
            }
        )
        return self._cached

    async def materialize_attachment_chunks(
        self,
        principal_id: str,
        conversation_id: str,
        *,
        content_sha256: str,
        parser_signature: str,
        chunk_signature: str,
        bundle: ParsedAttachmentBundle,
        ttl_days: int,
    ) -> bool:
        self.materialized.append(bundle)
        return True

    async def aiter_attachment_vectors(
        self,
        principal_id: str,
        conversation_id: str,
        references: list[tuple[int, attachments.AttachmentCacheKey]],
        *,
        ttl_days: int,
        page_size: int,
    ):
        self.vector_calls.append(
            {
                "principal_id": principal_id,
                "conversation_id": conversation_id,
                "references": references,
                "ttl_days": ttl_days,
                "page_size": page_size,
            }
        )
        if self._vector_rows:
            yield self._vector_rows

    async def aupdate_attachment_chunk_vectors(
        self,
        principal_id: str,
        conversation_id: str,
        chunks: list[AttachmentContextChunk],
        *,
        ttl_days: int,
    ) -> bool:
        self.updated.append(chunks)
        return True


def _dense_row(
    index: int,
    cache_key: attachments.AttachmentCacheKey,
) -> dict[str, Any]:
    return {
        "chunk_id": f"citation-{index}",
        "reference_id": "attachment",
        "content": ("evidence " * 30_000) if index == 0 else f"evidence {index}",
        "_cache_key": cache_key,
        "metadata": {},
    }


def _dense_embedder(*, query_vector: list[float], dimension: int = 2) -> Any:
    return SimpleNamespace(
        dimension=dimension,
        image_enabled=False,
        asymmetric=True,
        min_image_pixel=32,
        aembed_query=AsyncMock(return_value=query_vector),
        aembed_documents=AsyncMock(side_effect=AssertionError("documents must stay cached")),
    )


def _dense_service(
    *,
    test_config: Any,
    store: Any,
    embedder: Any,
    principal_id: str = "principal-1",
    conversation_id: str = "conversation-1",
) -> ComposerDocumentService:
    return ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
        principal_id=principal_id,
        conversation_id=conversation_id,
    )


async def test_dense_rankings_stream_ttl_scoped_pages_of_256(test_config: Any) -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    store = _SpyStore(
        None,
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                0,
                cache_key,
                signature,
                [1.0, 0.0],
            )
        ],
    )
    service = _dense_service(test_config=test_config, store=store, embedder=embedder)

    await service.adense_rankings("query", [_dense_row(0, cache_key)])

    assert store.vector_calls == [
        {
            "principal_id": "principal-1",
            "conversation_id": "conversation-1",
            "references": [(0, cache_key)],
            "ttl_days": 30,
            "page_size": 256,
        }
    ]


@pytest.mark.parametrize("database_vector", [None, [0.0, 1.0]])
async def test_dense_rankings_prefer_valid_request_local_vector_over_stale_or_missing_db(
    test_config: Any,
    database_vector: list[float] | None,
) -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    vector_rows = (
        [
            attachments.AttachmentVectorPageRow(
                0,
                cache_key,
                "stale-signature",
                database_vector,
            )
        ]
        if database_vector is not None
        else []
    )
    service = _dense_service(
        test_config=test_config,
        store=_SpyStore(None, vector_rows=vector_rows),
        embedder=embedder,
    )
    row = _dense_row(0, cache_key)

    ranked, trace = await service.adense_rankings(
        "query",
        [row],
        request_vectors={
            cache_key: attachments.AttachmentRequestVector(
                cache_key=cache_key,
                embedding_signature=signature,
                embedding_vector=[1.0, 0.0],
            )
        },
    )

    assert [item["chunk_id"] for item in ranked] == ["citation-0"]
    assert trace["composer_dense_status"] == "ranked"
    assert "embedding_signature" not in ranked[0]
    assert "embedding_vector" not in ranked[0]


async def test_dense_rankings_keep_valid_local_rows_when_pg_read_fails(
    test_config: Any,
) -> None:
    local_key = attachments.AttachmentCacheKey("local", "parser", "chunker", "local-cache")
    pg_key = attachments.AttachmentCacheKey("pg", "parser", "chunker", "pg-cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )

    class _FailingStore(_SpyStore):
        async def aiter_attachment_vectors(self, *args: Any, **kwargs: Any):
            self.vector_calls.append({"references": args[2]})
            raise RuntimeError("postgres unavailable")
            yield  # pragma: no cover - async generator shape

    store = _FailingStore(None)
    service = _dense_service(test_config=test_config, store=store, embedder=embedder)

    ranked, trace = await service.adense_rankings(
        "query",
        [_dense_row(0, local_key), _dense_row(1, pg_key)],
        request_vectors={
            local_key: attachments.AttachmentRequestVector(
                local_key,
                signature,
                [1.0, float(np.sqrt(np.float32(3.0)))],
            )
        },
    )

    assert [row["chunk_id"] for row in ranked] == ["citation-0"]
    assert trace == {
        "composer_dense_status": "ranked_degraded",
        "composer_dense_chunks": 1,
        "composer_dense_error": "RuntimeError",
    }
    assert store.vector_calls == [{"references": [(1, pg_key)]}]


async def test_dense_rankings_pg_failure_keeps_only_valid_local_candidates(
    test_config: Any,
) -> None:
    keys = [
        attachments.AttachmentCacheKey("content", "parser", "chunker", f"cache-{index}")
        for index in range(3)
    ]
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )

    class _FailingStore(_SpyStore):
        async def aiter_attachment_vectors(self, *args: Any, **kwargs: Any):
            raise RuntimeError("postgres unavailable")
            yield  # pragma: no cover - async generator shape

    service = _dense_service(
        test_config=test_config,
        store=_FailingStore(None),
        embedder=embedder,
    )

    ranked, trace = await service.adense_rankings(
        "query",
        [_dense_row(index, key) for index, key in enumerate(keys)],
        request_vectors={
            keys[0]: attachments.AttachmentRequestVector(keys[0], signature, [1.0, 0.0]),
            keys[1]: attachments.AttachmentRequestVector(keys[1], signature, [0.49, 0.872]),
        },
    )

    assert [row["chunk_id"] for row in ranked] == ["citation-0"]
    assert trace["composer_dense_status"] == "ranked_degraded"
    assert trace["composer_dense_error"] == "RuntimeError"


async def test_dense_rankings_pg_failure_without_local_candidates_reports_failed(
    test_config: Any,
) -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])

    class _FailingStore(_SpyStore):
        async def aiter_attachment_vectors(self, *args: Any, **kwargs: Any):
            raise RuntimeError("postgres unavailable")
            yield  # pragma: no cover - async generator shape

    service = _dense_service(
        test_config=test_config,
        store=_FailingStore(None),
        embedder=embedder,
    )

    ranked, trace = await service.adense_rankings(
        "query",
        [_dense_row(0, cache_key)],
    )

    assert ranked == []
    assert trace == {
        "composer_dense_status": "failed",
        "composer_dense_chunks": 0,
        "composer_dense_error": "RuntimeError",
    }


async def test_dense_rankings_trace_counts_all_admitted_rows(
    test_config: Any,
) -> None:
    keys = [
        attachments.AttachmentCacheKey("content", "parser", "chunker", f"cache-{index}")
        for index in range(3)
    ]
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )

    class _LaneStore(_SpyStore):
        async def aiter_attachment_vectors(
            self,
            principal_id: str,
            conversation_id: str,
            references: list[tuple[int, attachments.AttachmentCacheKey]],
            *,
            ttl_days: int,
            page_size: int,
        ):
            yield [
                attachments.AttachmentVectorPageRow(
                    global_order,
                    cache_key,
                    signature,
                    [1.0, 0.0],
                )
                for global_order, cache_key in references
            ]

    service = _dense_service(
        test_config=test_config,
        store=_LaneStore(None),
        embedder=embedder,
    )

    ranked, trace = await service.adense_rankings(
        "query",
        [
            _dense_row(0, keys[0]),
            _dense_row(1, keys[1]),
            {**_dense_row(2, keys[2]), "content": "history evidence " * 30_000},
        ],
    )

    assert [row["chunk_id"] for row in ranked] == ["citation-0", "citation-1", "citation-2"]
    assert trace == {
        "composer_dense_status": "ranked",
        "composer_dense_chunks": 3,
    }


async def test_dense_rankings_include_float32_half_and_ignore_invalid_vectors(
    test_config: Any,
) -> None:
    keys = [
        attachments.AttachmentCacheKey("content", "parser", "chunker", f"cache-{index}")
        for index in range(7)
    ]
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    store = _SpyStore(
        None,
        vector_rows=[
            attachments.AttachmentVectorPageRow(0, keys[0], signature, [1.0, 0.0]),
            attachments.AttachmentVectorPageRow(
                1,
                keys[1],
                signature,
                [1.0, float(np.sqrt(np.float32(3.0)))],
            ),
            attachments.AttachmentVectorPageRow(2, keys[2], signature, [0.49, 0.872]),
            attachments.AttachmentVectorPageRow(3, keys[3], signature, [0.0, 0.0]),
            attachments.AttachmentVectorPageRow(4, keys[4], signature, [float("nan"), 0.0]),
            attachments.AttachmentVectorPageRow(5, keys[5], signature, [1.0]),
            attachments.AttachmentVectorPageRow(6, keys[6], "stale", [1.0, 0.0]),
        ],
    )
    service = _dense_service(test_config=test_config, store=store, embedder=embedder)
    rows = [_dense_row(index, key) for index, key in enumerate(keys)]

    ranked, trace = await service.adense_rankings("query", rows)

    assert [row["chunk_id"] for row in ranked] == ["citation-0", "citation-1"]
    assert trace == {
        "composer_dense_status": "ranked",
        "composer_dense_chunks": 2,
    }


async def test_dense_rankings_accept_current_fused_provider_text_fallback(
    test_config: Any,
) -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    embedder.image_enabled = True
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
        fallback_reason="fused_provider_failed",
    )
    store = _SpyStore(
        None,
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                0,
                cache_key,
                signature,
                [1.0, 0.0],
            )
        ],
    )
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=True,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
        principal_id="principal-1",
        conversation_id="conversation-1",
    )
    row = {**_dense_row(0, cache_key), "image_data": "aW1hZ2U="}

    ranked, _ = await service.adense_rankings("query", [row])

    assert [item["chunk_id"] for item in ranked] == ["citation-0"]


async def test_dense_rankings_zero_query_disables_store_reads(test_config: Any) -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache")
    embedder = _dense_embedder(query_vector=[0.0, 0.0])
    store = _SpyStore(None)
    service = _dense_service(test_config=test_config, store=store, embedder=embedder)

    ranked, trace = await service.adense_rankings(
        "query",
        [_dense_row(0, cache_key)],
    )

    assert ranked == []
    assert trace == {
        "composer_dense_status": "no_query_vector",
        "composer_dense_chunks": 0,
    }
    assert store.vector_calls == []


async def test_dense_rankings_no_input_rows_skip_query_embedding(test_config: Any) -> None:
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    service = _dense_service(
        test_config=test_config,
        store=_SpyStore(None),
        embedder=embedder,
    )

    ranked, trace = await service.adense_rankings("query", [])

    assert ranked == []
    assert trace == {
        "composer_dense_status": "no_rows",
        "composer_dense_chunks": 0,
    }
    embedder.aembed_query.assert_not_awaited()


async def test_dense_rankings_no_valid_rows_report_zero_counts(test_config: Any) -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    store = _SpyStore(None)
    service = _dense_service(test_config=test_config, store=store, embedder=embedder)

    ranked, trace = await service.adense_rankings(
        "query",
        [_dense_row(0, cache_key)],
    )

    assert ranked == []
    assert trace == {
        "composer_dense_status": "no_rows",
        "composer_dense_chunks": 0,
    }
    embedder.aembed_query.assert_awaited_once_with("query")
    assert len(store.vector_calls) == 1


async def test_dense_rankings_rank_small_rows_passed_by_retrieval_router(test_config: Any) -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    store = _SpyStore(None)
    service = _dense_service(test_config=test_config, store=store, embedder=embedder)
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    row = {**_dense_row(0, cache_key), "content": "small evidence"}

    ranked, trace = await service.adense_rankings(
        "query",
        [row],
        request_vectors={
            cache_key: attachments.AttachmentRequestVector(
                cache_key,
                signature,
                [1.0, 0.0],
            )
        },
    )

    assert [item["chunk_id"] for item in ranked] == [row["chunk_id"]]
    assert trace["composer_dense_status"] == "ranked"
    embedder.aembed_query.assert_awaited_once_with("query")


async def test_dense_rankings_scan_all_rows_passed_by_retrieval_router(
    test_config: Any,
) -> None:
    short_key = attachments.AttachmentCacheKey("short", "parser", "chunker", "short-cache")
    long_key = attachments.AttachmentCacheKey("long", "parser", "chunker", "long-cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    store = _SpyStore(
        None,
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                global_order=0,
                cache_key=short_key,
                embedding_signature=signature,
                embedding_vector=[1.0, 0.0],
            ),
            attachments.AttachmentVectorPageRow(
                global_order=1,
                cache_key=long_key,
                embedding_signature=signature,
                embedding_vector=[1.0, 0.0],
            ),
        ],
    )
    service = _dense_service(test_config=test_config, store=store, embedder=embedder)
    short_row = {
        **_dense_row(1, short_key),
        "reference_id": "short-document",
        "content": "short evidence",
    }
    long_row = {
        **_dense_row(2, long_key),
        "reference_id": "long-document",
        "content": "long evidence " * 30_000,
    }

    ranked, trace = await service.adense_rankings(
        "query",
        [short_row, long_row],
    )

    assert [row["chunk_id"] for row in ranked] == ["citation-1", "citation-2"]
    assert trace["composer_dense_status"] == "ranked"
    embedder.aembed_query.assert_awaited_once_with("query")
    assert store.vector_calls[0]["references"] == [(0, short_key), (1, long_key)]


async def test_dense_rankings_propagate_query_cancellation(test_config: Any) -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    embedder.aembed_query.side_effect = asyncio.CancelledError
    service = _dense_service(
        test_config=test_config,
        store=_SpyStore(None),
        embedder=embedder,
    )

    with pytest.raises(asyncio.CancelledError):
        await service.adense_rankings("query", [_dense_row(0, cache_key)])


async def test_dense_rankings_match_one_shot_across_pages_and_preserve_duplicate_ties(
    test_config: Any,
) -> None:
    row_count = 513
    duplicate_key = attachments.AttachmentCacheKey(
        "duplicate-content",
        "parser",
        "chunker",
        "duplicate-cache",
    )
    keys = [
        duplicate_key
        if index in {255, 256}
        else attachments.AttachmentCacheKey(
            f"content-{index}",
            "parser",
            "chunker",
            f"cache-{index}",
        )
        for index in range(row_count)
    ]
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    source_vectors = np.asarray(
        [[1.0, float(index % 11) / 10.0] for index in range(row_count)],
        dtype=np.float32,
    )
    vectors_by_key = {
        cache_key: vector for cache_key, vector in zip(keys, source_vectors, strict=True)
    }

    class _PagedStore(_SpyStore):
        async def aiter_attachment_vectors(
            self,
            principal_id: str,
            conversation_id: str,
            references: list[tuple[int, attachments.AttachmentCacheKey]],
            *,
            ttl_days: int,
            page_size: int,
        ):
            self.vector_calls.append(
                {
                    "principal_id": principal_id,
                    "conversation_id": conversation_id,
                    "references": references,
                    "ttl_days": ttl_days,
                    "page_size": page_size,
                }
            )
            for start in range(0, len(references), page_size):
                yield [
                    attachments.AttachmentVectorPageRow(
                        global_order,
                        cache_key,
                        signature,
                        vectors_by_key[cache_key].tolist(),
                    )
                    for global_order, cache_key in references[start : start + page_size]
                ]

    store = _PagedStore(None)
    service = _dense_service(test_config=test_config, store=store, embedder=embedder)
    rows = [_dense_row(index, key) for index, key in enumerate(keys)]
    row_vectors = np.asarray([vectors_by_key[key] for key in keys], dtype=np.float32)
    expected_scores = row_vectors[:, 0] / np.linalg.norm(row_vectors, axis=1)
    assert np.array_equal(row_vectors[255], row_vectors[256])
    assert expected_scores[255] == expected_scores[256]
    expected_indices = sorted(
        (index for index, score in enumerate(expected_scores) if score >= np.float32(0.5)),
        key=lambda index: (-float(expected_scores[index]), index),
    )

    ranked, _ = await service.adense_rankings("query", rows)

    assert [row["chunk_id"] for row in ranked] == [
        f"citation-{index}" for index in expected_indices
    ]
    duplicate_positions = [
        index
        for index, row in enumerate(ranked)
        if row["chunk_id"] in {"citation-255", "citation-256"}
    ]
    assert [ranked[index]["chunk_id"] for index in duplicate_positions] == [
        "citation-255",
        "citation-256",
    ]
    assert duplicate_positions[1] == duplicate_positions[0] + 1
    assert store.vector_calls[0]["page_size"] == 256


async def test_dense_rankings_failure_returns_error_type_and_empty_rankings(
    test_config: Any,
) -> None:
    cache_key = attachments.AttachmentCacheKey("content", "parser", "chunker", "cache")
    embedder = _dense_embedder(query_vector=[1.0, 0.0])
    embedder.aembed_query.side_effect = RuntimeError("provider unavailable")
    service = _dense_service(
        test_config=test_config,
        store=_SpyStore(None),
        embedder=embedder,
    )

    ranked, trace = await service.adense_rankings(
        "query",
        [_dense_row(0, cache_key), _dense_row(1, cache_key)],
    )

    assert ranked == []
    assert trace == {
        "composer_dense_status": "failed",
        "composer_dense_chunks": 0,
        "composer_dense_error": "RuntimeError",
    }


def test_parse_owner_shim_persist_is_a_noop_holding_no_store() -> None:
    shim = _ParseOwnerShim()

    # The shim owns no storage handles; parsing through it cannot write rows.
    assert not hasattr(shim, "full_docs")
    assert not hasattr(shim, "doc_status")
    assert not hasattr(shim, "chunks_vdb")
    assert shim._resolve_source_file_for_parser("/tmp/x.pdf") == "/tmp/x.pdf"


async def test_parse_owner_shim_persist_captures_sidecar_only() -> None:
    shim = _ParseOwnerShim()

    result = await shim._persist_parsed_full_docs(
        "att-doc", {"content": "body", "sidecar_location": "loc://x"}
    )

    # No-op returns nothing and only remembers the sidecar location.
    assert result is None
    assert shim.sidecar_location == "loc://x"


async def test_parse_attachment_to_bundle_native_path_writes_no_store(
    tmp_path: Path,
) -> None:
    # A .txt routes to the local legacy engine: no MinerU, no network.
    lightrag = _FakeLightRAG()
    text = "First paragraph about ships.\n\nSecond paragraph about the sea."

    bundle = await parse_attachment_to_bundle(
        lightrag=lightrag,
        attachment_id="att-1",
        filename="notes.txt",
        document_bytes=text.encode("utf-8"),
        parser_rules="",
    )

    assert bundle.chunks, "native parse produced no chunks"
    assert all(chunk.attachment_id == "att-1" for chunk in bundle.chunks)
    assert "ships" in " ".join(chunk.content for chunk in bundle.chunks)
    assert bundle.parser_signature.startswith("legacy:")


async def test_parse_attachment_to_bundle_docx_native_path(tmp_path: Path) -> None:
    # A .docx routes to the native engine, which imports python-docx AND
    # defusedxml. Pin that path end-to-end: a missing parser dependency used to
    # be swallowed as an attachment_parse_error (0 chunks, document silently
    # dropped) instead of failing a test.
    from docx import Document

    doc = Document()
    doc.add_paragraph("Fractions worksheet about numerators and denominators.")
    doc.add_paragraph("Add the numerators; keep the denominator.")
    path = tmp_path / "frac.docx"
    doc.save(str(path))

    bundle = await parse_attachment_to_bundle(
        lightrag=_FakeLightRAG(),
        attachment_id="att-docx",
        filename="frac.docx",
        document_bytes=path.read_bytes(),
        parser_rules="docx:native-iteP,*:mineru-iteP",
    )

    assert bundle.chunks, "native docx parse produced no chunks"
    assert all(chunk.attachment_id == "att-docx" for chunk in bundle.chunks)
    assert "numerators" in " ".join(chunk.content for chunk in bundle.chunks).lower()
    assert bundle.parser_signature.startswith("native")


async def test_service_returns_cache_hit_without_parsing(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    cache_key = attachments.AttachmentCacheKey("deadbeef", "legacy:", "chunker-v1", "t1")
    cached = ParsedAttachmentBundle(
        chunks=[replace(text_chunk("t1"), cache_key=cache_key)],
        parser_signature="legacy:",
    )
    embedder: Any = SimpleNamespace(dimension=2, image_enabled=False, aembed_documents=AsyncMock())
    embedding_signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    store = _SpyStore(
        cached=cached,
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                0,
                cache_key,
                embedding_signature,
                [1.0, 0.0],
            )
        ],
    )

    def _boom(*_args: Any, **_kwargs: Any) -> None:  # pragma: no cover - must not run
        raise AssertionError("parse must not run on a cache hit")

    monkeypatch.setattr(attachments, "parse_attachment_document", _boom)

    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )
    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="deadbeef",
    )

    assert trace["attachment_parse_cache_hit"] is True
    assert [chunk.content for chunk in bundle.chunks] == ["alpha"]
    assert bundle.evidence_mode == "full"
    assert bundle.chunks[0].embedding_signature is None
    assert bundle.chunks[0].embedding_vector is None
    assert store.vector_calls == []
    assert trace["attachment_vector_cache_hits"] == 0
    assert trace["attachment_vector_cache_misses"] == 0
    assert store.materialized == []


async def test_cache_hit_skips_parse_analysis_and_document_embedding(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    cache_key = attachments.AttachmentCacheKey(
        "content-v1",
        "legacy:",
        "chunker-v1",
        "cache-1",
    )
    cached_chunk = AttachmentContextChunk(
        chunk_id="att:att-1:0001",
        attachment_id="att-1",
        filename="report.txt",
        chunk_index=1,
        content="cached enriched text",
        token_estimate=3,
        cache_key=cache_key,
    )
    cached = ParsedAttachmentBundle(chunks=[cached_chunk], parser_signature="legacy:")
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(side_effect=AssertionError("embedding must not run")),
    )
    expected_signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    store = _SpyStore(
        cached,
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                global_order=0,
                cache_key=cache_key,
                embedding_signature=expected_signature,
                embedding_vector=[1.0, 0.0],
            )
        ],
    )

    def _boom(*_args: Any, **_kwargs: Any) -> None:
        raise AssertionError("parse or analysis must not run")

    monkeypatch.setattr(attachments, "parse_attachment_document", _boom, raising=False)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _boom, raising=False)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    embedder.aembed_documents.assert_not_awaited()
    assert bundle.evidence_mode == "full"
    assert bundle.chunks[0].embedding_signature is None
    assert bundle.chunks[0].embedding_vector is None
    assert trace["attachment_parse_cache_hit"] is True
    assert trace["attachment_vector_cache_hits"] == 0
    assert trace["attachment_vector_cache_misses"] == 0
    assert store.vector_calls == []
    assert store.materialized == []
    assert store.updated == []


@pytest.mark.parametrize(
    ("outcome", "rendered", "sidecar_type", "expected_count"),
    [
        ("success", True, "drawing", 1),
        ("intentionally_disabled", False, "drawing", 0),
    ],
)
async def test_cache_hit_trace_uses_persisted_analysis_and_rendered_markers(
    monkeypatch: Any,
    test_config: Any,
    outcome: str,
    rendered: bool,
    sidecar_type: str,
    expected_count: int,
) -> None:
    monkeypatch.setenv(
        "VLM_PROCESS_ENABLE",
        "true" if outcome == "success" else "false",
    )
    cache_key = attachments.AttachmentCacheKey(
        "content-v1",
        "legacy:",
        "chunker-v1",
        "cache-1",
    )
    metadata: dict[str, Any] = {"_composer_analysis_outcome": outcome}
    if rendered:
        metadata["_composer_mm_rendered"] = True
    cached_chunk = AttachmentContextChunk(
        chunk_id="att:att-1:0001",
        attachment_id="att-1",
        filename="report.txt",
        chunk_index=1,
        content="cached text",
        sidecar_type=sidecar_type,
        metadata=metadata,
        cache_key=cache_key,
    )
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(side_effect=AssertionError("embedding must not run")),
    )
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    store = _SpyStore(
        ParsedAttachmentBundle(chunks=[cached_chunk], parser_signature="legacy:"),
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                0,
                cache_key,
                signature,
                [1.0, 0.0],
            )
        ],
    )
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config.model_copy(
            update={
                "parser_sidecars": test_config.parser_sidecars.model_copy(
                    update={
                        "vlm": test_config.parser_sidecars.vlm.model_copy(
                            update={"enabled": outcome != "success"}
                        )
                    }
                )
            }
        ),
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    analysis_signature = json.loads(store.load_calls[0]["chunk_signature"])["analysis_signature"]
    assert json.loads(analysis_signature)["enabled"] is (outcome == "success")
    assert trace["attachment_analysis_outcome"] == outcome
    assert trace["attachment_mm_chunk_count"] == expected_count
    public_metadata = bundle.chunks[0].to_context_row()["metadata"]
    assert "_composer_analysis_outcome" not in public_metadata
    assert "_composer_mm_rendered" not in public_metadata


async def test_cache_hit_embedding_signature_failure_is_attachment_scoped(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    cache_key = attachments.AttachmentCacheKey(
        "content-v1",
        "legacy:",
        "chunker-v1",
        "cache-1",
    )
    cached = ParsedAttachmentBundle(
        chunks=[replace(retrieval_text_chunk("text-1"), cache_key=cache_key)],
        parser_signature="legacy:",
    )
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(),
    )

    def _signature_failure(**_kwargs: Any) -> str:
        raise RuntimeError("embedding signature unavailable")

    monkeypatch.setattr(attachments, "build_composer_embedding_signature", _signature_failure)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=_SpyStore(cached),
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    assert [chunk.content for chunk in bundle.chunks] == [_RETRIEVAL_CONTENT]
    assert trace["attachment_parse_cache_hit"] is True
    assert trace["attachment_embedding_error"] == "RuntimeError"
    assert trace["attachment_vector_cache_hits"] == 0
    assert trace["attachment_vector_cache_misses"] == 1
    assert trace["attachment_vector_update_status"] == "skipped"
    embedder.aembed_documents.assert_not_awaited()


async def test_fresh_small_bundle_routes_once_without_document_embedding(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield (
            attachments.ParsedAttachmentDocument("small text", "/tmp/live", "legacy:"),
            "P",
            "legacy:",
        )

    async def _disabled(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[text_chunk("small-1")],
            parser_signature="legacy:",
        )

    policy_calls = 0

    def _route_once(_chunks: Any) -> str:
        nonlocal policy_calls
        policy_calls += 1
        return "full"

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(side_effect=AssertionError("full mode must not embed")),
    )
    store = _SpyStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _disabled)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    monkeypatch.setattr(
        attachments,
        "_resolve_attachment_evidence_mode",
        _route_once,
        raising=False,
    )
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert policy_calls == 1
    embedder.aembed_documents.assert_not_awaited()
    assert bundle.evidence_mode == "full"
    assert bundle.chunks[0].embedding_signature is None
    assert bundle.chunks[0].embedding_vector is None
    assert trace["attachment_vector_cache_hits"] == 0
    assert trace["attachment_vector_cache_misses"] == 0
    assert store.materialized == [bundle]


async def test_cached_small_bundle_routes_once_without_reading_stale_vectors(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    cache_key = attachments.AttachmentCacheKey(
        "content-v1",
        "legacy:",
        "chunker-v1",
        "small-1",
    )
    cached = ParsedAttachmentBundle(
        chunks=[
            replace(
                text_chunk("small-1"),
                cache_key=cache_key,
                embedding_signature="stale-signature",
                embedding_vector=[1.0, 0.0],
            )
        ],
        parser_signature="legacy:",
    )
    store = _SpyStore(
        cached,
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                0,
                cache_key,
                "stale-signature",
                [1.0, 0.0],
            )
        ],
    )
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(side_effect=AssertionError("full mode must not embed")),
    )
    original_policy = attachments._resolve_attachment_evidence_mode
    policy_calls = 0

    def _route_once(chunks: Any) -> str:
        nonlocal policy_calls
        policy_calls += 1
        return original_policy(chunks)

    monkeypatch.setattr(attachments, "_resolve_attachment_evidence_mode", _route_once)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    assert policy_calls == 1
    assert bundle.evidence_mode == "full"
    assert bundle.chunks[0].embedding_signature is None
    assert bundle.chunks[0].embedding_vector is None
    assert store.vector_calls == []
    embedder.aembed_documents.assert_not_awaited()
    assert trace["attachment_vector_cache_hits"] == 0
    assert trace["attachment_vector_cache_misses"] == 0


async def test_fresh_oversized_bundle_routes_once_and_embeds_before_materialize(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    events: list[str] = []
    oversized_content = "retrieval evidence " * 30_000

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield (
            attachments.ParsedAttachmentDocument("large text", "/tmp/live", "legacy:"),
            "P",
            "legacy:",
        )

    async def _disabled(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        events.append("render")
        return ParsedAttachmentBundle(
            chunks=[
                replace(
                    text_chunk("large-1"),
                    content=oversized_content,
                    token_estimate=estimate_tokens(oversized_content),
                )
            ],
            parser_signature="legacy:",
        )

    original_policy = attachments._resolve_attachment_evidence_mode
    policy_calls = 0

    def _route_once(chunks: Any) -> str:
        nonlocal policy_calls
        policy_calls += 1
        events.append("route")
        return original_policy(chunks)

    async def _embed(_items: Any):
        events.append("embed")
        return (
            [DocumentEmbeddingVector("large-1", [1.0, 0.0], "text")],
            DocumentEmbeddingTrace(0, 1, 0, 0),
        )

    class _OrderedStore(_SpyStore):
        async def materialize_attachment_chunks(self, *args: Any, **kwargs: Any) -> bool:
            events.append("materialize")
            return await super().materialize_attachment_chunks(*args, **kwargs)

    store = _OrderedStore(cached=None)
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(side_effect=_embed),
    )
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _disabled)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    monkeypatch.setattr(attachments, "_resolve_attachment_evidence_mode", _route_once)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, _trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert policy_calls == 1
    assert events == ["render", "route", "embed", "materialize"]
    assert bundle.evidence_mode == "retrieval"
    assert bundle.chunks[0].embedding_vector == [1.0, 0.0]


async def test_cached_oversized_bundle_routes_once_and_refreshes_stale_vectors(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )

    oversized_content = "cached retrieval evidence " * 30_000
    cache_key = attachments.AttachmentCacheKey(
        "content-v1",
        "legacy:",
        "chunker-v1",
        "large-1",
    )
    cached = ParsedAttachmentBundle(
        chunks=[
            AttachmentContextChunk(
                chunk_id="large-1",
                attachment_id="att-1",
                filename="report.txt",
                chunk_index=1,
                content=oversized_content,
                token_estimate=estimate_tokens(oversized_content),
                cache_key=cache_key,
            )
        ],
        parser_signature="legacy:",
    )
    store = _SpyStore(
        cached,
        vector_rows=[
            attachments.AttachmentVectorPageRow(0, cache_key, "stale-signature", [1.0, 0.0])
        ],
    )
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(
            return_value=(
                [DocumentEmbeddingVector("large-1", [0.0, 1.0], "text")],
                DocumentEmbeddingTrace(0, 1, 0, 0),
            )
        ),
    )
    original_policy = attachments._resolve_attachment_evidence_mode
    policy_calls = 0

    def _route_once(chunks: Any) -> str:
        nonlocal policy_calls
        policy_calls += 1
        return original_policy(chunks)

    monkeypatch.setattr(attachments, "_resolve_attachment_evidence_mode", _route_once)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    assert policy_calls == 1
    assert bundle.evidence_mode == "retrieval"
    assert len(store.vector_calls) == 1
    embedder.aembed_documents.assert_awaited_once()
    assert bundle.chunks[0].embedding_vector == [0.0, 1.0]
    assert trace["attachment_vector_cache_hits"] == 0
    assert trace["attachment_vector_cache_misses"] == 1


async def test_cache_miss_runs_parse_analysis_chunk_embedding_then_one_materialize(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    events: list[str] = []
    scope_alive = False

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        nonlocal scope_alive
        events.append("enter")
        scope_alive = True
        try:
            yield (
                attachments.ParsedAttachmentDocument(
                    content="ordinary parser text",
                    blocks_path="/tmp/live.blocks.jsonl",
                    parser_signature="legacy:",
                ),
                "iteP",
                "legacy:",
            )
        finally:
            scope_alive = False
            events.append("exit")

    async def _analyze(**_kwargs: Any) -> ComposerAnalysisResult:
        assert scope_alive
        events.append("analyze")
        return ComposerAnalysisResult(
            ComposerAnalysisOutcome.SUCCESS,
            1,
            mm_chunks=({"content": "analyzed chart", "chunk_order_index": 0},),
        )

    async def _render(*, multimodal_chunks: Any, **_kwargs: Any) -> ParsedAttachmentBundle:
        assert scope_alive
        assert multimodal_chunks == ({"content": "analyzed chart", "chunk_order_index": 0},)
        events.append("render")
        return ParsedAttachmentBundle(
            chunks=[
                AttachmentContextChunk(
                    chunk_id="cache-1",
                    attachment_id="att-1",
                    filename="report.pdf",
                    chunk_index=1,
                    content=_RETRIEVAL_CONTENT,
                    token_estimate=estimate_tokens(_RETRIEVAL_CONTENT),
                    sidecar_type="drawing",
                )
            ],
            parser_signature="legacy:",
        )

    async def _embed(_items: Any):
        assert not scope_alive
        events.append("embed")
        return (
            [DocumentEmbeddingVector("cache-1", [1.0, 0.0], "text")],
            DocumentEmbeddingTrace(
                fused=0,
                text=1,
                fused_to_text_fallback=1,
                failed=0,
                error_type="RuntimeError",
            ),
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(side_effect=_embed),
    )
    store = _SpyStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope, raising=False)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _analyze, raising=False)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert events == ["enter", "analyze", "render", "exit", "embed"]
    assert len(store.materialized) == 1
    assert store.materialized[0] is bundle
    assert bundle.chunks[0].embedding_vector == [1.0, 0.0]
    assert trace["attachment_analysis_outcome"] == "success"
    assert trace["attachment_mm_chunk_count"] == 1
    assert trace["attachment_embedding_text"] == 1
    assert trace["attachment_embedding_error"] == "RuntimeError"


async def _assert_analysis_outcome_controls_text_cache_materialization(
    monkeypatch: Any,
    test_config: Any,
    outcome: str,
    error_type: str | None,
    expected_materializations: int,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield (
            attachments.ParsedAttachmentDocument("ordinary text", "/tmp/live", "legacy:"),
            "iteP",
            "legacy:",
        )

    async def _analyze(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(
            ComposerAnalysisOutcome(outcome),
            0,
            error_type=error_type,
        )

    async def _render(*, multimodal_chunks: Any, **_kwargs: Any) -> ParsedAttachmentBundle:
        assert multimodal_chunks == ()
        chunk = retrieval_text_chunk("text-1")
        if outcome == "degraded":
            chunk = replace(chunk, image_bytes=b"parser-sidecar-image")
        return ParsedAttachmentBundle(chunks=[chunk], parser_signature="legacy:")

    async def _embed(items: list[Any]):
        if outcome == "degraded":
            assert items[0].image_bytes is None
        return (
            [DocumentEmbeddingVector("text-1", [1.0, 0.0], "text")],
            DocumentEmbeddingTrace(0, 1, 0, 0),
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=True,
        aembed_documents=AsyncMock(side_effect=_embed),
    )
    store = _SpyStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _analyze)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=True,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert bundle.chunks[0].content == _RETRIEVAL_CONTENT
    assert bundle.chunks[0].embedding_signature is not None
    assert len(store.materialized) == expected_materializations
    assert trace["attachment_analysis_outcome"] == outcome
    assert trace["attachment_analysis_error"] == error_type
    assert trace["attachment_cache_materialized"] is bool(expected_materializations)
    assert trace["attachment_cache_write_status"] == (
        "written" if expected_materializations else "skipped"
    )


async def test_analysis_disabled_caches_text_chunks_and_vectors(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    await _assert_analysis_outcome_controls_text_cache_materialization(
        monkeypatch,
        test_config,
        "intentionally_disabled",
        None,
        1,
    )


async def test_mm_chunk_trace_counts_rendered_list_not_ordinary_sidecar_provenance(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield attachments.ParsedAttachmentDocument("text", "/tmp/live", "legacy:"), "P", "legacy:"

    async def _disabled(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[replace(text_chunk("text-1"), sidecar_type="drawing")],
            parser_signature="legacy:",
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(
            return_value=(
                [DocumentEmbeddingVector("text-1", [1.0, 0.0], "text")],
                DocumentEmbeddingTrace(0, 1, 0, 0),
            )
        ),
    )
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _disabled)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    store = _SpyStore(cached=None)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    _bundle_result, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert trace["attachment_mm_chunk_count"] == 0
    assert store.materialized[0].chunks[0].metadata == {
        "_composer_analysis_outcome": "intentionally_disabled"
    }


async def test_analysis_degraded_returns_request_local_text_but_does_not_cache(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    await _assert_analysis_outcome_controls_text_cache_materialization(
        monkeypatch,
        test_config,
        "degraded",
        "VLMFailure",
        0,
    )


async def test_partial_degradation_keeps_successful_visual_text_embeds_and_ranks_locally(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    successful_content = "successful analyzed sibling " * 30_000
    rendered_chunk = {
        "content": successful_content,
        "chunk_order_index": 0,
        "sidecar": {"type": "drawing", "id": "successful-chart"},
    }

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield (
            attachments.ParsedAttachmentDocument("ordinary text", "/tmp/live", "legacy:"),
            "iP",
            "legacy:",
        )

    async def _analyze(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(
            ComposerAnalysisOutcome.DEGRADED,
            1,
            error_type="SidecarAnalysisFailure",
            mm_chunks=(rendered_chunk,),
        )

    async def _render(*, multimodal_chunks: Any, **_kwargs: Any) -> ParsedAttachmentBundle:
        assert multimodal_chunks == (rendered_chunk,)
        return ParsedAttachmentBundle(
            chunks=[
                text_chunk("ordinary-text"),
                AttachmentContextChunk(
                    chunk_id="successful-visual",
                    attachment_id="att-1",
                    filename="report.pdf",
                    chunk_index=2,
                    content=successful_content,
                    sidecar_type="drawing",
                    image_bytes=b"successful-image",
                    image_mime_type="image/png",
                ),
            ],
            parser_signature="legacy:",
        )

    async def _embed(items: list[Any]):
        assert [item.image_bytes for item in items] == [None, None]
        return (
            [
                DocumentEmbeddingVector("ordinary-text", [0.0, 1.0], "text"),
                DocumentEmbeddingVector("successful-visual", [1.0, 0.0], "text"),
            ],
            DocumentEmbeddingTrace(0, 2, 0, 0),
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=True,
        asymmetric=True,
        min_image_pixel=32,
        aembed_query=AsyncMock(return_value=[1.0, 0.0]),
        aembed_documents=AsyncMock(side_effect=_embed),
    )
    store = _SpyStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _analyze)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=True,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
        principal_id="p1",
        conversation_id="c1",
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    successful = bundle.chunks[1]
    assert [chunk.chunk_id for chunk in bundle.chunks] == [
        "ordinary-text",
        "successful-visual",
    ]
    assert successful.image_bytes == b"successful-image"
    assert "failed" not in successful.content
    assert store.materialized == []
    assert trace["attachment_analysis_outcome"] == "degraded"
    assert trace["attachment_analysis_error"] == "SidecarAnalysisFailure"
    assert trace["attachment_mm_chunk_count"] == 1
    assert trace["attachment_embedding_fused"] == 0
    assert trace["attachment_embedding_text"] == 2
    assert trace["attachment_cache_write_status"] == "skipped"
    assert successful.cache_key is not None
    assert successful.embedding_signature is not None
    assert successful.embedding_vector is not None
    row = successful.to_context_row()
    ranked, _dense_trace = await service.adense_rankings(
        "query",
        [row],
        request_vectors={
            successful.cache_key: attachments.AttachmentRequestVector(
                successful.cache_key,
                successful.embedding_signature,
                successful.embedding_vector,
            )
        },
    )
    assert [item["chunk_id"] for item in ranked] == ["successful-visual"]
    assert ranked[0]["reference_id"] == "att-1"
    assert ranked[0]["image_data"]


@pytest.mark.parametrize(
    ("update_result", "expected_update_status"),
    [
        (True, "written"),
        (False, "skipped"),
        (RuntimeError("update unavailable"), "error"),
    ],
)
async def test_embedding_model_change_reuses_parse_analysis_and_refreshes_vectors(
    monkeypatch: Any,
    test_config: Any,
    update_result: bool | Exception,
    expected_update_status: str,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )

    cache_key = attachments.AttachmentCacheKey(
        "content-v1",
        "legacy:",
        "chunker-v1",
        "cache-1",
    )
    cached = ParsedAttachmentBundle(
        chunks=[
            AttachmentContextChunk(
                chunk_id="att:att-1:0001",
                attachment_id="att-1",
                filename="report.txt",
                chunk_index=1,
                content=_RETRIEVAL_CONTENT,
                token_estimate=estimate_tokens(_RETRIEVAL_CONTENT),
                cache_key=cache_key,
            )
        ],
        parser_signature="legacy:",
    )

    class _VectorUpdateStore(_SpyStore):
        async def aupdate_attachment_chunk_vectors(
            self,
            principal_id: str,
            conversation_id: str,
            chunks: list[AttachmentContextChunk],
            *,
            ttl_days: int,
        ) -> bool:
            await super().aupdate_attachment_chunk_vectors(
                principal_id,
                conversation_id,
                chunks,
                ttl_days=ttl_days,
            )
            if isinstance(update_result, Exception):
                raise update_result
            return update_result

    store = _VectorUpdateStore(
        cached,
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                global_order=0,
                cache_key=cache_key,
                embedding_signature="old-model-signature",
                embedding_vector=[1.0, 0.0],
            )
        ],
    )
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        asymmetric=True,
        min_image_pixel=32,
        aembed_query=AsyncMock(return_value=[0.0, 1.0]),
        aembed_documents=AsyncMock(
            return_value=(
                [DocumentEmbeddingVector("cache-1", [0.0, 1.0], "text")],
                DocumentEmbeddingTrace(0, 1, 0, 0, error_type="ValueError"),
            )
        ),
    )

    def _boom(*_args: Any, **_kwargs: Any) -> None:
        raise AssertionError("embedding changes must not reparse or reanalyze")

    monkeypatch.setattr(attachments, "parse_attachment_document", _boom)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _boom)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config.model_copy(
            update={"embedding": test_config.embedding.model_copy(update={"model": "new-model"})}
        ),
        principal_id="p1",
        conversation_id="c1",
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    embedder.aembed_documents.assert_awaited_once()
    assert bundle.chunks[0].content == _RETRIEVAL_CONTENT
    assert bundle.chunks[0].embedding_vector == [0.0, 1.0]
    assert len(store.updated) == 1
    assert store.updated[0] == bundle.chunks
    assert store.materialized == []
    assert trace["attachment_parse_cache_hit"] is True
    assert trace["attachment_vector_cache_hits"] == 0
    assert trace["attachment_vector_cache_misses"] == 1
    assert trace["attachment_embedding_error"] == "ValueError"
    assert trace["attachment_vector_update_status"] == expected_update_status
    if isinstance(update_result, Exception):
        assert trace["attachment_vector_update_error"] == "RuntimeError"
        chunk = bundle.chunks[0]
        assert chunk.cache_key is not None
        assert chunk.embedding_signature is not None
        assert chunk.embedding_vector is not None
        ranked, dense_trace = await service.adense_rankings(
            "query",
            [{**chunk.to_context_row(), "content": "refreshed evidence " * 30_000}],
            request_vectors={
                chunk.cache_key: attachments.AttachmentRequestVector(
                    chunk.cache_key,
                    chunk.embedding_signature,
                    chunk.embedding_vector,
                )
            },
        )
        assert [row["chunk_id"] for row in ranked] == [chunk.chunk_id]
        assert dense_trace["composer_dense_status"] == "ranked"
        assert "embedding_signature" not in ranked[0]
        assert "embedding_vector" not in ranked[0]


async def test_analysis_exception_is_stage_specific_and_keeps_rendered_text(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield attachments.ParsedAttachmentDocument("body", "", "legacy:"), "iteP", "legacy:"

    async def _analysis_failure(**_kwargs: Any):
        raise LookupError("analysis failed")

    async def _render(*, multimodal_chunks: Any, **_kwargs: Any) -> ParsedAttachmentBundle:
        assert multimodal_chunks == ()
        return ParsedAttachmentBundle(chunks=[text_chunk("text-1")], parser_signature="legacy:")

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(
            return_value=(
                [DocumentEmbeddingVector("text-1", [1.0, 0.0], "text")],
                DocumentEmbeddingTrace(0, 1, 0, 0),
            )
        ),
    )
    store = _SpyStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _analysis_failure)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert [chunk.content for chunk in bundle.chunks] == ["alpha"]
    assert trace["attachment_analysis_outcome"] == "degraded"
    assert trace["attachment_analysis_error"] == "LookupError"
    assert trace["attachment_parser_error"] is None
    assert trace["attachment_rendering_error"] is None
    assert trace["attachment_cache_write_status"] == "skipped"


async def test_rendering_exception_is_not_mislabeled_as_parser_failure(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield attachments.ParsedAttachmentDocument("body", "", "legacy:"), "", "legacy:"

    async def _disabled(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render_failure(**_kwargs: Any) -> ParsedAttachmentBundle:
        raise OSError("render failed")

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(),
    )
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _disabled)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render_failure)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=_SpyStore(cached=None),
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert bundle.chunks == []
    assert trace["attachment_parse_error"] == "OSError"
    assert trace["attachment_parser_error"] is None
    assert trace["attachment_rendering_error"] == "OSError"
    embedder.aembed_documents.assert_not_awaited()


async def test_embedding_exception_is_stage_specific_and_keeps_chunks_for_retry(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield attachments.ParsedAttachmentDocument("body", "", "legacy:"), "", "legacy:"

    async def _disabled(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[retrieval_text_chunk("text-1")],
            parser_signature="legacy:",
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(side_effect=TimeoutError("provider timeout")),
    )
    store = _SpyStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _disabled)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert [chunk.content for chunk in bundle.chunks] == [_RETRIEVAL_CONTENT]
    assert bundle.chunks[0].embedding_signature is None
    assert trace["attachment_parse_error"] is None
    assert trace["attachment_embedding_error"] == "TimeoutError"
    assert trace["attachment_embedding_failed"] == 1
    assert trace["attachment_cache_write_status"] == "written"


async def test_mixed_vector_cache_reuses_text_row_and_retries_only_visual_as_fused(
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )

    text_key = attachments.AttachmentCacheKey(
        "content-v1",
        "legacy:",
        "chunker-v1",
        "cache-text",
    )
    visual_key = attachments.AttachmentCacheKey(
        "content-v1",
        "legacy:",
        "chunker-v1",
        "cache-visual",
    )
    cached = ParsedAttachmentBundle(
        chunks=[
            replace(retrieval_text_chunk("text-1"), cache_key=text_key),
            replace(retrieval_visual_chunk("visual-1"), cache_key=visual_key),
        ],
        parser_signature="legacy:",
    )
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=True,
        aembed_documents=AsyncMock(
            return_value=(
                [DocumentEmbeddingVector("cache-visual", [0.0, 1.0], "fused")],
                DocumentEmbeddingTrace(1, 0, 0, 0),
            )
        ),
    )
    text_signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )
    fused_signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="fused",
    )
    store = _SpyStore(
        cached,
        vector_rows=[
            attachments.AttachmentVectorPageRow(
                global_order=0,
                cache_key=text_key,
                embedding_signature=text_signature,
                embedding_vector=[1.0, 0.0],
            ),
            attachments.AttachmentVectorPageRow(
                global_order=1,
                cache_key=visual_key,
                embedding_signature=text_signature,
                embedding_vector=[1.0, 0.0],
            ),
        ],
    )
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=True,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    embedder.aembed_documents.assert_awaited_once()
    (embed_inputs,) = embedder.aembed_documents.await_args.args
    assert [item.key for item in embed_inputs] == ["cache-visual"]
    assert embed_inputs[0].image_bytes == b"png-bytes"
    assert len(store.updated) == 1
    assert [chunk.cache_key for chunk in store.updated[0]] == [visual_key]
    assert bundle.chunks[0].embedding_signature == text_signature
    assert bundle.chunks[0].embedding_vector is None
    assert bundle.chunks[1].embedding_signature == fused_signature
    assert bundle.chunks[1].embedding_vector == [0.0, 1.0]
    assert json.loads(bundle.chunks[0].embedding_signature or "{}")["mode"] == "text"
    assert json.loads(bundle.chunks[1].embedding_signature or "{}")["mode"] == "fused"
    assert trace["attachment_vector_cache_hits"] == 1
    assert trace["attachment_vector_cache_misses"] == 1
    assert trace["attachment_embedding_fused"] == 1


async def test_vector_cache_validation_releases_each_page_and_refreshes_only_misses(
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )

    tracker = {"live": 0, "max_live": 0, "page_size": 2}

    class _TrackedVector(list[float]):
        def __init__(self, values: list[float]) -> None:
            super().__init__(values)
            tracker["live"] += 1
            tracker["max_live"] = max(tracker["max_live"], tracker["live"])

        def __del__(self) -> None:
            tracker["live"] -= 1

    cache_keys = [
        attachments.AttachmentCacheKey(
            "content-v1",
            "legacy:",
            "chunker-v1",
            f"cache-{index}",
        )
        for index in range(5)
    ]
    cached = ParsedAttachmentBundle(
        chunks=[
            replace(retrieval_text_chunk(f"text-{index}"), cache_key=cache_key)
            for index, cache_key in enumerate(cache_keys)
        ],
        parser_signature="legacy:",
    )
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(
            return_value=(
                [
                    DocumentEmbeddingVector("cache-1", [0.0, 1.0], "text"),
                    DocumentEmbeddingVector("cache-3", [0.0, 1.0], "text"),
                ],
                DocumentEmbeddingTrace(0, 2, 0, 0),
            )
        ),
    )
    signature = attachments.build_composer_embedding_signature(
        config=test_config,
        embedder=embedder,
        mode="text",
    )

    class _PagedStore(_SpyStore):
        async def aiter_attachment_vectors(self, *_args: Any, **_kwargs: Any):
            rows = [
                (signature, [1.0, 0.0]),
                ("stale", [1.0, 0.0]),
                (signature, [1.0, 0.0]),
                (signature, [0.0, 0.0]),
                (signature, [1.0, 0.0]),
            ]
            for offset in range(0, len(rows), tracker["page_size"]):
                page = [
                    attachments.AttachmentVectorPageRow(
                        global_order=index,
                        cache_key=cache_keys[index],
                        embedding_signature=rows[index][0],
                        embedding_vector=_TrackedVector(rows[index][1]),
                    )
                    for index in range(offset, min(offset + tracker["page_size"], len(rows)))
                ]
                yield page
                del page

    store = _PagedStore(cached)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    (embed_inputs,) = embedder.aembed_documents.await_args.args
    assert [item.key for item in embed_inputs] == ["cache-1", "cache-3"]
    assert trace["attachment_vector_cache_hits"] == 3
    assert trace["attachment_vector_cache_misses"] == 2
    assert tracker["max_live"] <= tracker["page_size"]
    assert tracker["live"] == 0
    assert [bundle.chunks[index].embedding_signature for index in (0, 2, 4)] == [
        signature,
        signature,
        signature,
    ]
    assert [bundle.chunks[index].embedding_vector for index in (0, 2, 4)] == [
        None,
        None,
        None,
    ]
    assert [bundle.chunks[index].embedding_vector for index in (1, 3)] == [
        [0.0, 1.0],
        [0.0, 1.0],
    ]


async def test_visual_chunk_uses_fused_vector_when_capability_active(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield (
            attachments.ParsedAttachmentDocument("text", "/tmp/live", "legacy:"),
            "iteP",
            "legacy:",
        )

    async def _analyze(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.SUCCESS, 1)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[
                AttachmentContextChunk(
                    chunk_id="visual-1",
                    attachment_id="att-1",
                    filename="report.pdf",
                    chunk_index=1,
                    content=_RETRIEVAL_CONTENT,
                    token_estimate=estimate_tokens(_RETRIEVAL_CONTENT),
                    sidecar_type="drawing",
                    image_bytes=b"sidecar-image",
                    image_mime_type="image/png",
                )
            ],
            parser_signature="legacy:",
        )

    async def _embed(items: list[Any]):
        assert items[0].image_bytes == b"sidecar-image"
        return (
            [DocumentEmbeddingVector("visual-1", [1.0, 0.0], "fused")],
            DocumentEmbeddingTrace(1, 0, 0, 0),
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=True,
        aembed_documents=AsyncMock(side_effect=_embed),
    )
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _analyze)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=_SpyStore(cached=None),
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=True,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    embed_items = embedder.aembed_documents.await_args.args[0]
    assert embed_items[0].image_bytes == b"sidecar-image"
    assert json.loads(bundle.chunks[0].embedding_signature or "{}")["mode"] == "fused"
    assert trace["attachment_embedding_fused"] == 1
    assert trace["attachment_embedding_text"] == 0


async def test_fused_failure_caches_text_fallback_but_remains_fused_retryable(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    parse_count = 0

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        nonlocal parse_count
        parse_count += 1
        yield (
            attachments.ParsedAttachmentDocument("text", "/tmp/live", "legacy:"),
            "iteP",
            "legacy:",
        )

    async def _analyze(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.SUCCESS, 1)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[
                AttachmentContextChunk(
                    chunk_id="visual-1",
                    attachment_id="att-1",
                    filename="report.pdf",
                    chunk_index=1,
                    content=_RETRIEVAL_CONTENT,
                    token_estimate=estimate_tokens(_RETRIEVAL_CONTENT),
                    sidecar_type="drawing",
                    image_bytes=b"sidecar-image",
                    image_mime_type="image/png",
                )
            ],
            parser_signature="legacy:",
        )

    embed_calls = 0

    async def _embed(items: list[Any]):
        nonlocal embed_calls
        embed_calls += 1
        assert items[0].image_bytes == b"sidecar-image"
        if embed_calls == 1:
            return (
                [
                    DocumentEmbeddingVector(
                        "visual-1",
                        [1.0, 0.0],
                        "text",
                        fallback_reason="fused_provider_failed",
                    )
                ],
                DocumentEmbeddingTrace(0, 1, 1, 0),
            )
        return (
            [DocumentEmbeddingVector("visual-1", [0.0, 1.0], "fused")],
            DocumentEmbeddingTrace(1, 0, 0, 0),
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=True,
        aembed_documents=AsyncMock(side_effect=_embed),
    )
    store = _SpyStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _analyze)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=True,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config.model_copy(
            update={
                "parser_sidecars": test_config.parser_sidecars.model_copy(
                    update={
                        "vlm": test_config.parser_sidecars.vlm.model_copy(update={"enabled": True})
                    }
                )
            }
        ),
    )

    first_bundle, first_trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )
    fallback_chunk = first_bundle.chunks[0]
    fallback_signature = json.loads(fallback_chunk.embedding_signature or "{}")
    assert fallback_signature["mode"] == "text"
    assert fallback_signature["fallback_reason"] == "fused_provider_failed"
    assert first_trace["attachment_embedding_fallback"] == 1
    assert fallback_chunk.cache_key is not None
    store._cached = replace(
        first_bundle,
        chunks=[replace(fallback_chunk, embedding_signature=None, embedding_vector=None)],
    )
    store._vector_rows = [
        attachments.AttachmentVectorPageRow(
            global_order=0,
            cache_key=fallback_chunk.cache_key,
            embedding_signature=fallback_chunk.embedding_signature,
            embedding_vector=fallback_chunk.embedding_vector,
        )
    ]

    second_bundle, second_trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    assert parse_count == 1
    assert embed_calls == 2
    assert len(store.materialized) == 1
    assert len(store.updated) == 1
    assert json.loads(second_bundle.chunks[0].embedding_signature or "{}")["mode"] == "fused"
    assert second_trace["attachment_parse_cache_hit"] is True
    assert second_trace["attachment_analysis_outcome"] == "success"
    assert second_trace["attachment_vector_cache_misses"] == 1
    assert second_trace["attachment_embedding_fused"] == 1


@pytest.mark.parametrize(
    "image_bytes",
    [b"not an image", _png_bytes(size=(2, 2))],
    ids=["corrupt", "undersized"],
)
async def test_deterministic_image_rejection_reuses_cached_text_without_reopening(
    monkeypatch: Any,
    test_config: Any,
    image_bytes: bytes,
) -> None:
    import dlightrag.core.document_embedding as document_embedding
    from dlightrag.core.document_embedding import RobustDocumentEmbedder
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield (
            attachments.ParsedAttachmentDocument("text", "/tmp/live", "legacy:"),
            "iteP",
            "legacy:",
        )

    async def _analyze(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.SUCCESS, 1)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[
                AttachmentContextChunk(
                    chunk_id="visual-1",
                    attachment_id="att-1",
                    filename="report.pdf",
                    chunk_index=1,
                    content=_RETRIEVAL_CONTENT,
                    token_estimate=estimate_tokens(_RETRIEVAL_CONTENT),
                    sidecar_type="drawing",
                    image_bytes=image_bytes,
                    image_mime_type="image/png",
                )
            ],
            parser_signature="legacy:",
        )

    provider: Any = SimpleNamespace(
        asymmetric=False,
        embed_index_fused=AsyncMock(),
        embed_texts=AsyncMock(return_value=[[1.0, 0.0]]),
    )
    embedder = RobustDocumentEmbedder(
        embedder=provider,
        image_enabled=True,
        dimension=2,
        min_image_pixel=4,
        batch_size=1,
        max_concurrency=1,
    )
    open_calls = 0
    original_open = document_embedding._open_valid_image

    def _tracked_open(*args: Any, **kwargs: Any):
        nonlocal open_calls
        open_calls += 1
        return original_open(*args, **kwargs)

    monkeypatch.setattr(document_embedding, "_open_valid_image", _tracked_open)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _analyze)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    store = _SpyStore(cached=None)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=True,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config.model_copy(
            update={
                "parser_sidecars": test_config.parser_sidecars.model_copy(
                    update={
                        "vlm": test_config.parser_sidecars.vlm.model_copy(update={"enabled": True})
                    }
                )
            }
        ),
    )

    first_bundle, _first_trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )
    rejected_chunk = first_bundle.chunks[0]
    rejected_signature = json.loads(rejected_chunk.embedding_signature or "{}")
    assert rejected_signature["mode"] == "text"
    assert rejected_signature["fallback_reason"] == "image_rejected"
    assert rejected_chunk.cache_key is not None
    store._cached = replace(
        first_bundle,
        chunks=[replace(rejected_chunk, embedding_signature=None, embedding_vector=None)],
    )
    store._vector_rows = [
        attachments.AttachmentVectorPageRow(
            global_order=0,
            cache_key=rejected_chunk.cache_key,
            embedding_signature=rejected_chunk.embedding_signature,
            embedding_vector=rejected_chunk.embedding_vector,
        )
    ]

    second_bundle, second_trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"ignored",
        content_sha256="content-v1",
    )

    assert open_calls == 1
    provider.embed_index_fused.assert_not_awaited()
    assert provider.embed_texts.await_count == 1
    assert second_bundle.chunks[0].embedding_signature == rejected_chunk.embedding_signature
    assert second_bundle.chunks[0].embedding_vector is None
    assert second_trace["attachment_vector_cache_hits"] == 1
    assert second_trace["attachment_vector_cache_misses"] == 0


async def test_parser_cancellation_reraises_cancelled_error_releases_tempdir_and_writes_no_cache(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    cleaned = False

    @asynccontextmanager
    async def _cancelled_parse(**_kwargs: Any):
        nonlocal cleaned
        try:
            raise asyncio.CancelledError("parser cancelled")
            yield  # pragma: no cover
        finally:
            cleaned = True

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(),
    )
    store = _SpyStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _cancelled_parse)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    with pytest.raises(asyncio.CancelledError, match="parser cancelled"):
        await service.achunks_for_attachment(
            principal_id="p1",
            conversation_id="c1",
            attachment_id="att-1",
            filename="report.txt",
            document_bytes=b"body",
            content_sha256="content-v1",
        )

    assert cleaned is True
    assert store.materialized == []
    assert store.updated == []
    embedder.aembed_documents.assert_not_awaited()


async def test_parse_source_write_and_tempdir_cleanup_run_off_event_loop(
    monkeypatch: Any,
) -> None:
    from lightrag.parser import registry

    event_loop_thread = threading.get_ident()
    write_threads: list[int] = []
    cleanup_threads: list[int] = []
    original_write_bytes = Path.write_bytes
    original_temporary_directory = attachments.tempfile.TemporaryDirectory

    def _tracked_write_bytes(path: Path, data: bytes) -> int:
        write_threads.append(threading.get_ident())
        return original_write_bytes(path, data)

    class _TrackedTemporaryDirectory:
        def __init__(self, **kwargs: Any) -> None:
            self._inner = original_temporary_directory(**kwargs)
            self.name = self._inner.name

        def __enter__(self) -> str:
            return self.name

        def __exit__(self, *_args: Any) -> None:
            self.cleanup()

        def cleanup(self) -> None:
            cleanup_threads.append(threading.get_ident())
            self._inner.cleanup()

    class _Parser:
        async def parse(self, _ctx: Any) -> Any:
            return SimpleNamespace(content="body", blocks_path="")

    monkeypatch.setattr(Path, "write_bytes", _tracked_write_bytes)
    monkeypatch.setattr(attachments.tempfile, "TemporaryDirectory", _TrackedTemporaryDirectory)
    monkeypatch.setattr(registry, "get_parser", lambda _engine: _Parser())

    async with attachments.parse_attachment_document(
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        parser_rules="",
    ):
        pass

    assert write_threads and all(thread != event_loop_thread for thread in write_threads)
    assert cleanup_threads and all(thread != event_loop_thread for thread in cleanup_threads)


async def test_real_parse_cancellation_waits_for_tempdir_cleanup(monkeypatch: Any) -> None:
    from lightrag.parser import registry

    parser_started = asyncio.Event()
    source_path: Path | None = None

    class _Parser:
        async def parse(self, ctx: Any) -> Any:
            nonlocal source_path
            source_path = Path(ctx.file_path)
            parser_started.set()
            await asyncio.Event().wait()

    monkeypatch.setattr(registry, "get_parser", lambda _engine: _Parser())

    async def _parse() -> None:
        async with attachments.parse_attachment_document(
            attachment_id="att-1",
            filename="report.txt",
            document_bytes=b"body",
            parser_rules="",
        ):
            pass

    task = asyncio.create_task(_parse())
    await asyncio.wait_for(parser_started.wait(), timeout=1)
    task.cancel("parser cancelled")
    with pytest.raises(asyncio.CancelledError, match="parser cancelled"):
        await task

    assert source_path is not None
    assert not source_path.parent.exists()


async def test_embedding_cancellation_reraises_cancelled_error_and_writes_no_partial_materialization(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    cleaned = False

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        nonlocal cleaned
        try:
            yield (
                attachments.ParsedAttachmentDocument("text", "/tmp/live", "legacy:"),
                "iteP",
                "legacy:",
            )
        finally:
            cleaned = True

    async def _analyze(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[retrieval_text_chunk("text-1")],
            parser_signature="legacy:",
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(side_effect=asyncio.CancelledError("embedding cancelled")),
    )
    store = _SpyStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _analyze)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    with pytest.raises(asyncio.CancelledError, match="embedding cancelled"):
        await service.achunks_for_attachment(
            principal_id="p1",
            conversation_id="c1",
            attachment_id="att-1",
            filename="report.txt",
            document_bytes=b"body",
            content_sha256="content-v1",
        )

    assert cleaned is True
    assert store.materialized == []
    assert store.updated == []


async def test_parse_context_keeps_mutated_sidecars_alive_through_rendering(
    monkeypatch: Any,
    test_config: Any,
    tmp_path: Path,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    observed_paths: list[Path] = []

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        scope = tmp_path / "live-scope"
        scope.mkdir()
        blocks = scope / "report.blocks.jsonl"
        sidecar = scope / "report.drawings.json"
        image = scope / "chart.png"
        blocks.write_text("", encoding="utf-8")
        image.write_bytes(b"image")
        sidecar.write_text(
            json.dumps({"drawings": {"chart": {"path": "chart.png"}}}),
            encoding="utf-8",
        )
        try:
            yield (
                attachments.ParsedAttachmentDocument("body", str(blocks), "legacy:"),
                "iteP",
                "legacy:",
            )
        finally:
            for path in scope.iterdir():
                path.unlink()
            scope.rmdir()

    async def _analyze(*, parsed_data: dict[str, Any], **_kwargs: Any) -> ComposerAnalysisResult:
        blocks = Path(parsed_data["blocks_path"])
        sidecar = blocks.with_name("report.drawings.json")
        payload = json.loads(sidecar.read_text(encoding="utf-8"))
        payload["drawings"]["chart"]["llm_analyze_result"] = {
            "status": "success",
            "text": "mutated chart description",
        }
        sidecar.write_text(json.dumps(payload), encoding="utf-8")
        observed_paths.extend([blocks, sidecar])
        return ComposerAnalysisResult(ComposerAnalysisOutcome.SUCCESS, 1)

    async def _render(*, parsed: Any, **_kwargs: Any) -> ParsedAttachmentBundle:
        sidecar = Path(parsed.blocks_path).with_name("report.drawings.json")
        payload = json.loads(sidecar.read_text(encoding="utf-8"))
        description = payload["drawings"]["chart"]["llm_analyze_result"]["text"]
        assert description == "mutated chart description"
        return ParsedAttachmentBundle(
            chunks=[
                AttachmentContextChunk(
                    chunk_id="visual-1",
                    attachment_id="att-1",
                    filename="report.pdf",
                    chunk_index=1,
                    content=description,
                    sidecar_type="drawing",
                    image_bytes=b"image",
                )
            ],
            parser_signature="legacy:",
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
    )

    async def _embed(_items: list[Any]):
        assert observed_paths
        assert all(not path.exists() for path in observed_paths)
        return (
            [DocumentEmbeddingVector("visual-1", [1.0, 0.0], "text")],
            DocumentEmbeddingTrace(0, 1, 0, 0),
        )

    embedder.aembed_documents = AsyncMock(side_effect=_embed)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _analyze)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=_SpyStore(cached=None),
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, _trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert bundle.chunks[0].content == "mutated chart description"
    assert observed_paths
    assert all(not path.exists() for path in observed_paths)


async def test_identical_bytes_in_another_conversation_do_not_hit_cache(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    class _ConversationStore(_SpyStore):
        async def load_attachment_chunks(
            self, principal_id: str, conversation_id: str, *args: Any, **kwargs: Any
        ):
            await super().load_attachment_chunks(principal_id, conversation_id, *args, **kwargs)
            return self._cached if conversation_id == "c1" else None

    parse_count = 0

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        nonlocal parse_count
        parse_count += 1
        yield attachments.ParsedAttachmentDocument("body", "", "legacy:"), "", "legacy:"

    async def _disabled(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[retrieval_text_chunk("text-1")],
            parser_signature="legacy:",
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(
            return_value=(
                [DocumentEmbeddingVector("text-1", [1.0, 0.0], "text")],
                DocumentEmbeddingTrace(0, 1, 0, 0),
            )
        ),
    )
    store = _ConversationStore(cached=ParsedAttachmentBundle(chunks=[]))
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _disabled)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    _bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c2",
        attachment_id="att-2",
        filename="report.txt",
        document_bytes=b"same-bytes",
        content_sha256="same-sha",
    )

    assert parse_count == 1
    assert trace["attachment_parse_cache_hit"] is False
    assert store.load_calls[0]["conversation_id"] == "c2"


async def test_expired_conversation_never_materializes_derived_results(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    class _ExpiredStore(_SpyStore):
        async def materialize_attachment_chunks(self, *args: Any, **kwargs: Any) -> bool:
            await super().materialize_attachment_chunks(*args, **kwargs)
            self.materialized.clear()
            return False

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield attachments.ParsedAttachmentDocument("body", "", "legacy:"), "", "legacy:"

    async def _disabled(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[retrieval_text_chunk("text-1")],
            parser_signature="legacy:",
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(
            return_value=(
                [DocumentEmbeddingVector("text-1", [1.0, 0.0], "text")],
                DocumentEmbeddingTrace(0, 1, 0, 0),
            )
        ),
    )
    store = _ExpiredStore(cached=None)
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _disabled)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=7,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="expired",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert bundle.chunks[0].embedding_vector == [1.0, 0.0]
    assert trace["attachment_cache_materialized"] is False
    assert trace["attachment_cache_write_status"] == "skipped"
    assert trace["attachment_cache_write_error"] is None
    assert store.materialized == []
    assert store.load_calls[0]["ttl_days"] == 7


async def test_cache_io_failures_keep_request_local_enriched_bundle(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    class _FailingCacheStore(_SpyStore):
        async def load_attachment_chunks(self, *args: Any, **kwargs: Any):
            raise RuntimeError("cache read unavailable")

        async def materialize_attachment_chunks(self, *args: Any, **kwargs: Any) -> bool:
            raise ValueError("cache write unavailable")

    @asynccontextmanager
    async def _parse_scope(**_kwargs: Any):
        yield attachments.ParsedAttachmentDocument("body", "", "legacy:"), "", "legacy:"

    async def _disabled(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return ParsedAttachmentBundle(
            chunks=[retrieval_text_chunk("text-1")],
            parser_signature="legacy:",
        )

    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        asymmetric=True,
        min_image_pixel=32,
        aembed_query=AsyncMock(return_value=[1.0, 0.0]),
        aembed_documents=AsyncMock(
            return_value=(
                [DocumentEmbeddingVector("text-1", [1.0, 0.0], "text")],
                DocumentEmbeddingTrace(0, 1, 0, 0),
            )
        ),
    )
    monkeypatch.setattr(attachments, "parse_attachment_document", _parse_scope)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _disabled)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=_FailingCacheStore(cached=None),
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
        principal_id="p1",
        conversation_id="c1",
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="content-v1",
    )

    assert bundle.chunks[0].embedding_vector == [1.0, 0.0]
    assert trace["attachment_cache_error"] == "RuntimeError"
    assert trace["attachment_cache_read_error"] == "RuntimeError"
    assert trace["attachment_cache_write_error"] == "ValueError"
    assert trace["attachment_cache_write_status"] == "error"
    assert trace["attachment_cache_materialized"] is False
    chunk = bundle.chunks[0]
    assert chunk.cache_key is not None
    assert chunk.embedding_signature is not None
    assert chunk.embedding_vector is not None
    row = {
        **chunk.to_context_row(),
        "content": "request-local evidence " * 30_000,
    }
    ranked, dense_trace = await service.adense_rankings(
        "query",
        [row],
        request_vectors={
            chunk.cache_key: attachments.AttachmentRequestVector(
                cache_key=chunk.cache_key,
                embedding_signature=chunk.embedding_signature,
                embedding_vector=chunk.embedding_vector,
            )
        },
    )
    assert [item["chunk_id"] for item in ranked] == [chunk.chunk_id]
    assert dense_trace["composer_dense_status"] == "ranked"
    assert "embedding_signature" not in ranked[0]
    assert "embedding_vector" not in ranked[0]


async def test_service_cache_miss_parses_and_materializes(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    from dlightrag.core.document_embedding import (
        DocumentEmbeddingTrace,
        DocumentEmbeddingVector,
    )
    from dlightrag.core.request.composer_analysis import (
        ComposerAnalysisOutcome,
        ComposerAnalysisResult,
    )

    store = _SpyStore(cached=None)
    parsed = ParsedAttachmentBundle(chunks=[text_chunk("t1")], parser_signature="legacy:")

    @asynccontextmanager
    async def _fake_parse(**_kwargs: Any):
        yield attachments.ParsedAttachmentDocument("body", "", "legacy:"), "", "legacy:"

    async def _disabled(**_kwargs: Any) -> ComposerAnalysisResult:
        return ComposerAnalysisResult(ComposerAnalysisOutcome.INTENTIONALLY_DISABLED, 0)

    async def _render(**_kwargs: Any) -> ParsedAttachmentBundle:
        return parsed

    monkeypatch.setattr(attachments, "parse_attachment_document", _fake_parse)
    monkeypatch.setattr(attachments, "aanalyze_composer_sidecars", _disabled)
    monkeypatch.setattr(attachments, "build_attachment_bundle_from_parse_result", _render)
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(
            return_value=(
                [DocumentEmbeddingVector("t1", [1.0, 0.0], "text")],
                DocumentEmbeddingTrace(0, 1, 0, 0),
            )
        ),
    )

    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )
    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="deadbeef",
    )

    assert trace["attachment_parse_cache_hit"] is False
    assert [chunk.content for chunk in bundle.chunks] == ["alpha"]
    assert store.materialized == [bundle]


async def test_service_parse_failure_is_attachment_scoped(
    monkeypatch: Any,
    test_config: Any,
) -> None:
    store = _SpyStore(cached=None)

    @asynccontextmanager
    async def _fail(**_kwargs: Any):
        raise RuntimeError("mineru exploded")
        yield  # pragma: no cover

    monkeypatch.setattr(attachments, "parse_attachment_document", _fail)
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(),
    )

    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )
    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-1",
        filename="report.txt",
        document_bytes=b"body",
        content_sha256="deadbeef",
    )

    assert bundle.chunks == []
    assert trace["attachment_parse_error"] == "RuntimeError"
    assert trace["attachment_parser_error"] == "RuntimeError"
    assert trace["attachment_directive_error"] is None
    assert trace["attachment_parse_cache_hit"] is False
    assert store.materialized == []


async def test_invalid_parser_hint_is_attachment_scoped(test_config: Any) -> None:
    store = _SpyStore(cached=None)
    embedder: Any = SimpleNamespace(
        dimension=2,
        image_enabled=False,
        aembed_documents=AsyncMock(),
    )
    service = ComposerDocumentService(
        lightrag=_FakeLightRAG(),
        store=store,
        parser_rules="",
        ttl_days=30,
        robust_document_embedder=embedder,
        direct_image_embedding_enabled=False,
        model_bundle=SimpleNamespace(vlm_identity={}, extract_identity={}),
        config=test_config,
    )

    bundle, trace = await service.achunks_for_attachment(
        principal_id="p1",
        conversation_id="c1",
        attachment_id="att-invalid",
        filename="report.[unknown].txt",
        document_bytes=b"body",
        content_sha256="deadbeef",
    )

    assert bundle.chunks == []
    assert trace["attachment_parse_error"] == "FilenameParserHintError"
    assert trace["attachment_directive_error"] == "FilenameParserHintError"
    assert trace["attachment_parser_error"] is None
    assert trace["attachment_analysis_outcome"] == "degraded"
    assert trace["attachment_parse_cache_hit"] is False
    assert store.load_calls == []
    assert store.materialized == []


def test_parser_and_chunk_signatures_are_stable_strings() -> None:
    lightrag = _FakeLightRAG()
    parser_sig = resolve_attachment_parser_signature("report.txt", "")
    chunk_sig = resolve_attachment_chunk_signature(lightrag, "report.txt", "")

    assert parser_sig == resolve_attachment_parser_signature("report.txt", "")
    assert chunk_sig == resolve_attachment_chunk_signature(lightrag, "report.txt", "")
    assert parser_sig.startswith("legacy:")
    assert "tokenizer" in chunk_sig


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(pytest.main([__file__, "-q"]))
